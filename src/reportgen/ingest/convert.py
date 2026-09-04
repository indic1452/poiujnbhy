"""Конвертация исходных файлов библиотеки в Markdown.

Слой приёма превращает то, что реально лежит у инженеров на диске (PDF,
DOCX, заметки в .md и .txt), в единый Markdown, который дальше нарезается на
чанки функциями из :mod:`reportgen.corpus`. Формат чанка при этом не меняется —
конвертер лишь готовит для него текст.

Что важно знать про результат конвертации:

* **Заголовки.** В PDF их нет как сущности, есть только кегль шрифта, поэтому
  заголовки восстанавливаются по размеру шрифта относительно основного текста
  (медиана по всему документу). В DOCX заголовки берутся из стилей ``Heading N``.
* **Номера страниц.** Для PDF в текст вставляются невидимые в разметке маркеры
  ``<!-- page: N -->``. По ним слой приёма проставляет ``meta['page']`` каждому
  чанку — без этого ссылка в отчёте («с. 47») невозможна.
* **Сканы.** Если текста в PDF почти нет, документ помечается ``needs_ocr``:
  система не притворяется, что разобрала скан, а честно говорит, что нужен OCR.
* **Битые файлы не роняют приём.** Любая ошибка разбора превращается в
  предупреждение и пустой текст: один испорченный файл не должен останавливать
  индексацию каталога на тысячу документов.
"""

from __future__ import annotations

import os

import hashlib
import re
import unicodedata
import shutil
import tempfile
from contextlib import contextmanager
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterator, List, Sequence, Tuple

from ..packages import pip_hint
from ..corpus import DOC_TYPES

__all__ = [
    "ConvertedDocument",
    "MissingDependencyError",
    "SUPPORTED_SUFFIXES",
    "BUILTIN_SUFFIXES",
    "read_text",
    "decode_bytes",
    "clean_line",
    "external_path",
    "supported_suffixes",
    "format_support",
    "convert_file",
    "first_page_number",
    "guess_doc_type",
    "page_marker",
    "page_markers",
    "sha256_file",
    "strip_page_markers",
]

#: Расширения, которые разбирает само ядро, без модулей форматов и внешних
#: программ. Полный список того, что система читает на этой машине, отдаёт
#: supported_suffixes(); он же доступен как SUPPORTED_SUFFIXES (см. __getattr__).
BUILTIN_SUFFIXES = (".pdf", ".docx", ".dotx", ".md", ".markdown", ".txt")
DEFAULT_DOC_TYPE = "literature"

#: Ниже этого числа символов на страницу PDF считается сканом (нужен OCR).
MIN_CHARS_PER_PAGE = 100

#: Ниже этой доли пробелов текст считается склеенным. В русском и английском
#: тексте пробел приходится примерно на каждый шестой-седьмой знак (доля
#: 0.13–0.17); в таблицах и листингах — реже, поэтому порог взят с большим
#: запасом. Склеенный текст даёт из абзаца одно слово: по нему не находит ни
#: поиск словами (в указателе одно длинное «слово»), ни смысловой поиск.
MIN_SPACE_SHARE = 0.04

#: Короткие куски не проверяем: в подписи «Рис.1» пробелов и не должно быть.
GLUE_CHECK_CHARS = 400


def _repairs_note(repairs: Dict[str, int]) -> str:
    """Одна строка о том, что система поправила в тексте файла.

    Человеку важно не «текст поправлен», а что именно с ним сделали: правок
    много — документ стоит пересохранить у себя, а не жить с починкой.
    """
    names = {
        "ligatures": ("лигатура", "лигатуры", "лигатур"),
        "invisible": ("невидимый знак — мягкий перенос или нулевой пробел",
                      "невидимых знака", "невидимых знаков"),
        "scripts": ("степень или индекс записаны знаками ^ и _",
                    "степени и индекса записаны знаками ^ и _",
                    "степеней и индексов записаны знаками ^ и _"),
        "homoglyphs": ("слово с латинскими буквами внутри русского",
                       "слова с латинскими буквами внутри русских",
                       "слов с латинскими буквами внутри русских"),
    }
    parts = []
    for key, count in repairs.items():
        if not count:
            continue
        forms = names.get(key)
        word = _plural(count, *forms) if forms else key
        parts.append(f"{count} {word}")
    return ("текст файла поправлен при чтении: "
            + ", ".join(parts)
            + " — иначе такие слова не находятся поиском")


def _plural(count: int, one: str, few: str, many: str) -> str:
    """Согласование числа со словом по-русски: 1 лигатура, 2 лигатуры, 5 лигатур."""
    tail = abs(int(count)) % 100
    if 11 <= tail <= 14:
        return many
    tail %= 10
    if tail == 1:
        return one
    if 2 <= tail <= 4:
        return few
    return many


def glued_text_warning(text: str) -> str:
    """«Методыцифровогокодирования» — так выглядит плохо разобранный PDF.

    Беда тихая: документ в библиотеке есть, страниц много, фрагментов много,
    а найти в нём нельзя ничего. Инженер видит документ в списке и считает,
    что он работает.
    """
    body = "".join(str(text or "").split("\n"))
    if len(body) < GLUE_CHECK_CHARS:
        return ""
    letters = sum(1 for ch in body if ch.isalpha())
    if letters < GLUE_CHECK_CHARS // 2:
        return ""
    spaces = body.count(" ")
    share = spaces / len(body)
    if share >= MIN_SPACE_SHARE:
        return ""
    return (f"текст склеен без пробелов (пробелов {share * 100:.1f}% при "
            f"обычных 13–17%) — файл разобран плохо, поиск по такому "
            f"документу почти ничего не найдёт; переложите его через OCR "
            f"или сохраните в другом формате")


#: Во сколько раз строка должна быть крупнее основного текста, чтобы считаться
#: заголовком. 1.12 — компромисс: подзаголовки обычно крупнее тела на 15–30 %.
HEADING_MIN_RATIO = 1.12
#: Длинный абзац крупным шрифтом (например, аннотация) заголовком не считается.
HEADING_MAX_CHARS = 200
#: Больше трёх уровней в восстановленной структуре смысла не имеет.
MAX_HEADING_LEVEL = 3

_READ_BLOCK = 1 << 20

#: Маркер страницы. Это HTML-комментарий: в разметке он не виден, в индекс не
#: попадает (слой приёма вырезает его), но позволяет привязать чанк к странице.
_PAGE_MARKER_RE = re.compile(r"<!--\s*page:\s*(\d+)\s*-->")

_HYPHENS = "-‐‑­−"
_HEADING_STYLE_RE = re.compile(r"^(?:heading|заголовок)\s*(\d+)$", re.IGNORECASE)
_TITLE_STYLES = {"title", "название", "заголовок"}
_SUBTITLE_STYLES = {"subtitle", "подзаголовок"}
_CAPTION_STYLES = {"caption", "название объекта", "подпись", "подрисуночная подпись"}
_LIST_STYLE_MARKERS = ("list", "список", "bullet", "маркированный", "нумерованный")

_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


from .text_repair import (                # noqa: E402 — рядом со своими помощниками
    drop_running_titles,
    repair_report,
    repair_text,
)


class MissingDependencyError(RuntimeError):
    """Не установлен пакет, нужный для разбора формата."""


@dataclass
class ConvertedDocument:
    """Результат конвертации одного файла в Markdown."""

    text: str = ""
    title: str = ""
    page_count: int = 0
    meta: Dict[str, Any] = field(default_factory=dict)
    needs_ocr: bool = False
    warnings: List[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        """Пусто ли содержимое (с учётом того, что маркеры страниц — не текст)."""
        return not strip_page_markers(self.text).strip()

    @property
    def char_count(self) -> int:
        return len(strip_page_markers(self.text).strip())

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "page_count": self.page_count,
            "chars": self.char_count,
            "needs_ocr": self.needs_ocr,
            "meta": self.meta,
            "warnings": list(self.warnings),
        }


# ------------------------------------------------------------ утилиты ----

def page_marker(page: int) -> str:
    """Маркер начала страницы, вставляемый в Markdown."""
    return f"<!-- page: {int(page)} -->"


def strip_page_markers(text: str) -> str:
    """Убирает маркеры страниц и лишние пустые строки после них."""
    cleaned = _PAGE_MARKER_RE.sub("", text)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def page_markers(text: str) -> List[Tuple[int, int]]:
    """Все маркеры страниц во фрагменте: пары (номер страницы, смещение в тексте)."""
    return [(int(match.group(1)), match.start()) for match in _PAGE_MARKER_RE.finditer(text)]


def first_page_number(text: str) -> int | None:
    """Номер первой страницы, упомянутой в фрагменте, или ``None``."""
    match = _PAGE_MARKER_RE.search(text)
    return int(match.group(1)) if match else None


def sha256_file(path: str | Path) -> str:
    """SHA-256 файла. Читается блоками: даташит на 300 МБ в память не влезет."""
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(_READ_BLOCK), b""):
            digest.update(block)
    return digest.hexdigest()


def guess_doc_type(path: str | Path, root: str | Path | None = None,
                   default: str | None = DEFAULT_DOC_TYPE) -> str | None:
    """Тип документа по имени каталога верхнего уровня внутри корпуса.

    Правило то же, что в :func:`reportgen.corpus.load_file`: ``standards/ГОСТ.pdf``
    даёт ``standards``. Если каталог не опознан (или файл лежит прямо в корне
    корпуса), возвращается ``default``.

    ``default=None`` означает «каталог ничего не сказал» — по этому признаку
    приём понимает, что тип надо определять по содержимому: библиотеку часто
    приносят как есть, папкой «Разное» или выгрузкой со старого сервера.
    """
    path = Path(path)
    if root is not None:
        try:
            relative = _resolve(path).relative_to(_resolve(Path(root)))
        except ValueError:
            relative = None
        if relative is not None:
            if len(relative.parts) > 1 and relative.parts[0].lower() in DOC_TYPES:
                return relative.parts[0].lower()
            return default
    for parent in _resolve(path).parents:
        if parent.name.lower() in DOC_TYPES:
            return parent.name.lower()
    return default


@contextmanager
def external_path(path: Path, suffix: str | None = None) -> Iterator[Path]:
    """Путь к файлу, безопасный для сторонних программ.

    Часть внешних инструментов (djvulibre, а в некоторых сборках tesseract и
    LibreOffice) не открывает файлы, в имени или пути которых есть кириллица:
    они собраны без поддержки юникода в аргументах и молча возвращают ошибку
    «файл не найден». Проверено: cjb2 на файле «скан.pbm» падает, на «scan.pbm»
    работает.

    В библиотеке отдела по-русски названо почти всё, поэтому перед вызовом
    внешней программы файл при необходимости копируется во временный каталог
    под латинским именем. Копия удаляется автоматически.
    """
    text = str(path)
    if text.isascii():
        yield path
        return
    with tempfile.TemporaryDirectory(prefix="reportgen-tool-") as directory:
        target = Path(directory) / f"document{suffix or path.suffix.lower()}"
        shutil.copyfile(path, target)
        yield target


def _resolve(path: Path) -> Path:
    try:
        return path.resolve()
    except OSError:  # pragma: no cover — экзотические ФС
        return path.absolute()


def _reason(error: BaseException) -> str:
    text = str(error).strip()
    return text or error.__class__.__name__


def _join_wrapped(left: str, right: str) -> str:
    """Склеивает строку, разорванную переносом слова."""
    if (
        len(left) >= 2
        and left[-1] in _HYPHENS
        and left[-2].isalpha()
        and right[:1].isalpha()
        and right[:1].islower()
    ):
        return left[:-1] + right
    return f"{left} {right}"


def clean_line(text: str) -> str:
    """Нормализует строку: неразрывные пробелы, мягкие переносы, хвостовые пробелы."""
    text = text.replace("­", "").replace(" ", " ").replace("﻿", "")
    text = "".join(
        character
        for character in text
        if character in "\t\n" or not unicodedata.category(character).startswith("C")
    )
    return re.sub(r"[ \t]{2,}", " ", text).strip()


def _weighted_median(samples: Sequence[Tuple[float, int]]) -> float:
    """Медиана размеров шрифта, взвешенная по числу символов.

    Взвешивание принципиально: колонтитулы и номера страниц набраны мелко, но
    их много, а по числу символов основной текст всё равно перевешивает.
    """
    pairs = sorted((float(size), max(1, int(weight))) for size, weight in samples)
    total = sum(weight for _, weight in pairs)
    if not total:
        return 0.0
    half = total / 2
    accumulated = 0
    for size, weight in pairs:
        accumulated += weight
        if accumulated >= half:
            return size
    return pairs[-1][0]


# ------------------------------------------------------- ленивый импорт ---

def _import_pymupdf():
    try:
        import pymupdf  # type: ignore
    except ImportError as error:  # pragma: no cover — зависит от окружения
        raise MissingDependencyError(
            "для разбора PDF нужен пакет pymupdf (%s); " % pip_hint("pymupdf") +
            "установите его в изолированном контуре из локального зеркала"
        ) from error
    _silence_mupdf(pymupdf)
    return pymupdf


def _reset_mupdf_log(pymupdf) -> None:
    """Очистить накопленные сообщения MuPDF перед разбором файла."""
    tools = getattr(pymupdf, "TOOLS", None)
    reset = getattr(tools, "reset_mupdf_warnings", None) if tools else None
    if callable(reset):
        try:
            reset()
        except Exception:  # noqa: BLE001
            pass


def _mupdf_complaints(pymupdf) -> int:
    """Сколько раз MuPDF пожаловался на разметку при разборе файла."""
    tools = getattr(pymupdf, "TOOLS", None)
    getter = getattr(tools, "mupdf_warnings", None) if tools else None
    if not callable(getter):
        return 0
    try:
        logged = getter()
    except Exception:  # noqa: BLE001
        return 0
    if not logged:
        return 0
    if isinstance(logged, str):
        return len([line for line in logged.splitlines() if line.strip()])
    return len(logged)


def _silence_mupdf(pymupdf) -> None:
    """Убрать поток «MuPDF error: syntax error» из консоли.

    Библиотека MuPDF пишет о каждой шероховатости разметки PDF прямо в
    стандартный поток ошибок, минуя Python. На типографском файле, собранном
    старым генератором, это сотни строк вида «syntax error in content stream»
    и «unknown keyword: 'Tj21EB0A091j'» — при том, что MuPDF после каждой
    успешно продолжает, и текст извлекается полностью.

    Инженер видит экран, залитый красными строками, и решает, что приём
    сломался. Настоящие сообщения — сколько файлов принято и что не
    разобралось — в этом потоке теряются.

    Сообщения не выбрасываются: они копятся внутри MuPDF, и разборщик PDF
    забирает их числом («разметка PDF повреждена в N местах, текст извлечён»).
    """
    tools = getattr(pymupdf, "TOOLS", None)
    if tools is None:  # pragma: no cover — очень старая сборка
        return
    for name in ("mupdf_display_errors", "mupdf_display_warnings"):
        setter = getattr(tools, name, None)
        if callable(setter):
            try:
                setter(False)
            except Exception:  # noqa: BLE001 — молчание не стоит падения приёма
                pass


def _import_docx():
    try:
        import docx  # type: ignore
    except ImportError as error:  # pragma: no cover — зависит от окружения
        raise MissingDependencyError(
            "для разбора DOCX нужен пакет python-docx (%s); " % pip_hint("python-docx") +
            "установите его в изолированном контуре из локального зеркала"
        ) from error
    return docx


# ------------------------------------------------------------------ PDF ---

#: Промежуток между спанами, начиная с которого между ними был пробел.
#: Считается в долях кегля: пробел в наборных шрифтах — от четверти до трети
#: круглой. Порог ниже трети, потому что вёрстка пробелы поджимает.
SPAN_GAP_SHARE = 0.22

#: Спан считается индексом или степенью, если он мельче строки во столько
#: раз. Обычные выделения (полужирный, курсив) кегля не меняют, а индекс в
#: наборе всегда заметно мельче.
SCRIPT_SIZE_SHARE = 0.82

#: И если его базовая линия сдвинута хотя бы на такую долю кегля. Без
#: проверки сдвига под правило попала бы любая мелкая сноска в строке.
SCRIPT_SHIFT_SHARE = 0.14


def _span_text(span: Any) -> str:
    return str(span.get("text", "") or "")


def _line_text(spans: Sequence[Any]) -> Tuple[str, float]:
    """Строка из спанов: с пробелами там, где они были, и с индексами.

    Раньше спаны просто склеивались подряд. В PDF спан обрывается на каждой
    смене шрифта, и абзац, где часть слов полужирные, приходил из файла как
    «Занимаемаяполосачастотизмеряетсяметодом» — документ ложился в библиотеку
    целым на вид и не находился ни по одному слову.

    Пробела в файле может не быть вовсе: вёрстка ставит слово на своё место
    координатой, а не пробелом. Поэтому смотрим не на текст, а на разрыв
    между правым краем предыдущего спана и левым краем следующего.

    Тем же взглядом видно степень и индекс: спан мельче строки и приподнят
    или опущен относительно её базовой линии. Такой спан записываем как
    «^2» и «_вх» — иначе «4πR2» читается как число 2 при R, а не как
    квадрат, и в ответе помощника оказывается не то, что в стандарте.
    """
    body = [span for span in spans if _span_text(span).strip()]
    if not body:
        return "", 0.0
    sizes = [(float(span.get("size", 0.0) or 0.0), len(_span_text(span)))
             for span in body]
    line_size = _weighted_median(sizes) or max(size for size, _ in sizes)
    base = _weighted_median([
        (float((span.get("origin") or (0.0, 0.0))[1]), len(_span_text(span)))
        for span in body]) or 0.0

    parts: List[str] = []
    mode = ""                       # какой ряд идёт сейчас: '^', '_' или никакой
    previous_right: float | None = None
    for span in body:
        text = _span_text(span)
        size = float(span.get("size", 0.0) or 0.0)
        box = span.get("bbox") or (0.0, 0.0, 0.0, 0.0)
        left, right = float(box[0]), float(box[2])
        origin_y = float((span.get("origin") or (0.0, base))[1])

        shift = base - origin_y          # больше нуля — спан приподнят
        script = ""
        if (line_size and size and size <= line_size * SCRIPT_SIZE_SHARE
                and abs(shift) >= line_size * SCRIPT_SHIFT_SHARE):
            script = "^" if shift > 0 else "_"

        gap = 0.0 if previous_right is None else left - previous_right
        need_space = (
            previous_right is not None
            and gap >= max(line_size, size) * SPAN_GAP_SHARE
            and not parts[-1].endswith((" ", "-"))
            and not text.startswith(" ")
        )

        if script != mode:
            # Закрываем прежний ряд и открываем новый. Разрыв между двумя
            # индексами подряд («P_вх» и «P_изл») не должен их слепить.
            mode = script
            if script:
                if need_space:
                    parts.append(" ")
                parts.append(script)
                need_space = False
        if need_space:
            parts.append(" ")
        parts.append(text)
        previous_right = right

    return clean_line("".join(parts)), line_size


def _blocks_in_reading_order(raw_blocks: Sequence[Any], width: float) -> List[Any]:
    """Блоки по порядку чтения, с учётом двух колонок.

    MuPDF сортирует блоки сверху вниз, и на двухколоночной странице книги
    строки левой и правой колонки чередуются: получается текст, в котором
    предложения перебивают друг друга. Ни поиск, ни модель такого не читают.

    Колонки ищем по одному признаку — есть ли по середине страницы полоса,
    которую не пересекает ни один блок. Есть — читаем сначала левую колонку
    целиком, потом правую. Нет — оставляем как было: страница одноколоночная
    или свёрстана сложнее, и угадывать тут нельзя.
    """
    boxes = [(block, block.get("bbox") or (0.0, 0.0, 0.0, 0.0))
             for block in raw_blocks]
    if len(boxes) < 4 or width <= 0:
        return list(raw_blocks)
    middle = width / 2
    for _, box in boxes:
        if float(box[0]) < middle < float(box[2]):
            return list(raw_blocks)       # блок пересекает середину — не колонки
    left = [(block, box) for block, box in boxes if float(box[2]) <= middle]
    right = [(block, box) for block, box in boxes if float(box[0]) >= middle]
    if not left or not right:
        return list(raw_blocks)
    order = (sorted(left, key=lambda item: float(item[1][1]))
             + sorted(right, key=lambda item: float(item[1][1])))
    return [block for block, _ in order]


def _pdf_page_blocks(page: Any) -> List[List[Tuple[str, float]]]:
    """Блоки страницы: список абзацев, абзац — список пар (строка, кегль)."""
    data = page.get_text("dict", sort=True)
    width = float(data.get("width") or 0.0)
    blocks: List[List[Tuple[str, float]]] = []
    text_blocks = [raw for raw in data.get("blocks", [])
                   if raw.get("type", 0) == 0]  # тип 1 — картинка, текста нет
    for raw_block in _blocks_in_reading_order(text_blocks, width):
        lines: List[Tuple[str, float]] = []
        for raw_line in raw_block.get("lines", []):
            text, size = _line_text(raw_line.get("spans", []))
            if not text:
                continue
            lines.append((text, size))
        if lines:
            blocks.append(lines)
    return blocks


def _heading_levels(blocks: Sequence[Sequence[Tuple[str, float]]], body_size: float) -> Dict[float, int]:
    """Сопоставляет кегли заголовков уровням '#', '##', '###'.

    Абсолютные пороги («в 1.35 раза крупнее — это h1») ломаются на первом же
    документе с другой вёрсткой, поэтому кегли ранжируются: самый крупный из
    встреченных заголовочных размеров становится первым уровнем.
    """
    if body_size <= 0:
        return {}
    candidates = set()
    for block in blocks:
        text = " ".join(line for line, _ in block)
        if len(text) > HEADING_MAX_CHARS:
            continue
        size = max(size for _, size in block)
        if size >= body_size * HEADING_MIN_RATIO:
            candidates.add(round(size * 2) / 2)
    ranked = sorted(candidates, reverse=True)[:MAX_HEADING_LEVEL]
    return {size: level for level, size in enumerate(ranked, start=1)}


def _render_blocks(
    blocks: Sequence[Sequence[Tuple[str, float]]], levels: Dict[float, int]
) -> List[str]:
    """Превращает блоки страницы в куски Markdown."""
    parts: List[Tuple[str, str]] = []  # (вид, текст)
    for block in blocks:
        text = ""
        for line, _ in block:
            text = line if not text else _join_wrapped(text, line)
        if not text:
            continue
        size = round(max(size for _, size in block) * 2) / 2
        level = levels.get(size) if len(text) <= HEADING_MAX_CHARS else None
        if level:
            parts.append(("heading", f"{'#' * level} {text}"))
            continue
        if parts and parts[-1][0] == "paragraph" and parts[-1][1].endswith(tuple(_HYPHENS)):
            # Абзац разорван между блоками (частый случай в двухколоночной
            # вёрстке): слово с переносом склеиваем обратно.
            parts[-1] = ("paragraph", _join_wrapped(parts[-1][1], text))
            continue
        parts.append(("paragraph", text))
    return [text for _, text in parts]


def _convert_pdf(path: Path) -> ConvertedDocument:
    """PDF → Markdown через PyMuPDF, с восстановлением заголовков по кеглю."""
    result = ConvertedDocument(title=path.stem, meta={"source_format": "pdf"})
    try:
        pymupdf = _import_pymupdf()
    except MissingDependencyError as error:
        result.warnings.append(str(error))
        return result

    _reset_mupdf_log(pymupdf)
    try:
        document = pymupdf.open(str(path))
    except Exception as error:
        result.warnings.append(f"не удалось открыть PDF: {_reason(error)}")
        return result

    try:
        if getattr(document, "needs_pass", False):
            result.warnings.append(
                "PDF защищён паролем — текст не извлечён; расшифруйте файл перед приёмом"
            )
            return result
        result.page_count = int(getattr(document, "page_count", 0) or 0)
        result.meta["page_count"] = result.page_count

        pages: List[List[List[Tuple[str, float]]]] = []
        for number in range(result.page_count):
            try:
                pages.append(_pdf_page_blocks(document[number]))
            except Exception as error:
                result.warnings.append(
                    f"страница {number + 1}: текст не извлечён ({_reason(error)})"
                )
                pages.append([])

        metadata = dict(getattr(document, "metadata", None) or {})
    except Exception as error:
        result.warnings.append(f"PDF разобран не полностью: {_reason(error)}")
        return result
    finally:
        try:
            document.close()
        except Exception:  # pragma: no cover — закрытие уже закрытого документа
            pass

    # MuPDF ругается на каждую шероховатость разметки, но после каждой
    # успешно продолжает. Раньше это выливалось в консоль сотнями строк
    # «syntax error in content stream» и топило настоящий вывод приёма.
    # Теперь — одна строка в карточке документа, и та справочная.
    complaints = _mupdf_complaints(pymupdf)
    if complaints:
        result.meta["pdf_repairs"] = complaints
        result.warnings.append(
            f"разметка PDF повреждена в {complaints} местах — MuPDF восстановил, "
            "текст извлечён; сверьте выборочно числа и обозначения"
        )

    samples = [
        (size, len(line))
        for page in pages
        for block in page
        for line, size in block
    ]
    body_size = _weighted_median(samples)
    levels = _heading_levels([block for page in pages for block in page], body_size)
    result.meta["body_font_size"] = round(body_size, 2)

    rendered_pages = [_render_blocks(page, levels) for page in pages]
    # Колонтитул «2 специальный отдел — Методика измерений 17» на каждой из
    # шестисот страниц книги попадал в каждый фрагмент: смысл фрагмента
    # разбавлялся названием отдела, а поиск по названию отдела находил всю
    # библиотеку целиком.
    rendered_pages, running = drop_running_titles(rendered_pages)
    if running:
        result.meta["running_titles"] = running
    pieces: List[str] = []
    for number, rendered in enumerate(rendered_pages, start=1):
        if not rendered:
            continue
        pieces.append(page_marker(number))
        pieces.extend(rendered)
    raw_text = "\n\n".join(pieces)
    result.text = repair_text(raw_text)
    repairs = repair_report(raw_text, result.text)
    if repairs:
        result.meta["text_repairs"] = repairs
        result.warnings.append(_repairs_note(repairs))

    characters = sum(length for _, length in samples)
    if result.page_count and characters < MIN_CHARS_PER_PAGE * result.page_count:
        result.needs_ocr = True
        per_page = characters / result.page_count
        result.warnings.append(
            f"извлечено всего {characters} символов ({per_page:.0f} на страницу) — "
            "похоже на скан; нужен слой OCR (PaddleOCR или tesseract с языками rus+eng), "
            "документ в индекс не берём"
        )
    if not result.page_count:
        result.warnings.append("в PDF нет ни одной страницы")
    if not result.needs_ocr:
        # Проверяем по тому, что получилось, а не по тому, что было в файле:
        # пробелы между словами система теперь восстанавливает сама, и
        # жаловаться надо лишь на то, чего восстановить не удалось.
        glued = glued_text_warning(strip_page_markers(result.text))
        if glued:
            result.warnings.append(glued)

    title = str(metadata.get("title") or "").strip()
    if not title:
        title = _first_markdown_heading(result.text)
    result.title = title or path.stem
    for key in ("author", "subject", "keywords"):
        value = str(metadata.get(key) or "").strip()
        if value:
            result.meta[key] = value
    return result


# ----------------------------------------------------------------- DOCX ---

def _style_name(paragraph: Any) -> str:
    try:
        return (paragraph.style.name or "").strip()
    except Exception:  # pragma: no cover — документ без таблицы стилей
        return ""


def _docx_heading_level(style: str) -> int | None:
    lowered = style.strip().lower()
    match = _HEADING_STYLE_RE.match(lowered)
    if match:
        return min(int(match.group(1)), 6)
    if lowered in _TITLE_STYLES:
        return 1
    if lowered in _SUBTITLE_STYLES:
        return 2
    return None


def _is_list_style(style: str) -> bool:
    lowered = style.strip().lower()
    if lowered in _CAPTION_STYLES:
        return False
    return any(marker in lowered for marker in _LIST_STYLE_MARKERS)


def _cell_text(cell: Any) -> str:
    lines = [clean_line(paragraph.text) for paragraph in cell.paragraphs]
    text = " <br> ".join(line for line in lines if line)
    return text.replace("|", "\\|").strip()


def _table_to_markdown(table: Any) -> str:
    """Таблица DOCX → таблица Markdown (целиком, без разрезания на части)."""
    rows: List[List[str]] = []
    for row in table.rows:
        cells: List[str] = []
        seen: set[int] = set()
        for cell in row.cells:
            # Объединённые ячейки python-docx возвращает несколько раз подряд.
            marker = id(cell._tc)
            if marker in seen:
                continue
            seen.add(marker)
            cells.append(_cell_text(cell))
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _paragraph_images(element: Any) -> int:
    return sum(
        1
        for child in element.iter()
        if child.tag in (f"{_WORD_NS}drawing", f"{_WORD_NS}pict")
    )


def _convert_docx(path: Path) -> ConvertedDocument:
    """DOCX → Markdown: стили заголовков, таблицы, списки, подписи к рисункам."""
    result = ConvertedDocument(title=path.stem, meta={"source_format": "docx"})
    try:
        docx = _import_docx()
    except MissingDependencyError as error:
        result.warnings.append(str(error))
        return result

    try:
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore

        document = docx.Document(str(path))
    except Exception as error:
        result.warnings.append(f"не удалось открыть DOCX: {_reason(error)}")
        return result

    pieces: List[str] = []
    heading_title = ""
    images = 0
    try:
        for child in document.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "tbl":
                try:
                    markdown = _table_to_markdown(Table(child, document))
                except Exception as error:
                    result.warnings.append(f"таблица пропущена: {_reason(error)}")
                    continue
                if markdown:
                    pieces.append(markdown)
                continue
            if tag != "p":
                continue
            paragraph = Paragraph(child, document)
            text = clean_line(paragraph.text)
            count = _paragraph_images(child)
            if count:
                for number in range(images + 1, images + count + 1):
                    # Картинки не извлекаем, но место в тексте помечаем: инженер
                    # должен видеть, что здесь был рисунок (док. 03, шаг 1).
                    pieces.append(f"![рисунок {number}]()")
                images += count
            if not text:
                continue
            style = _style_name(paragraph)
            level = _docx_heading_level(style)
            if level:
                pieces.append(f"{'#' * level} {text}")
                if not heading_title and level == 1:
                    heading_title = text
                continue
            if style.strip().lower() in _CAPTION_STYLES:
                # Подпись к рисунку/таблице — обычный текст: она несёт смысл и
                # должна попадать в поиск вместе с окружающим разделом.
                pieces.append(text)
                continue
            if _is_list_style(style):
                pieces.append(f"- {text}")
                continue
            pieces.append(text)
    except Exception as error:
        result.warnings.append(f"DOCX разобран не полностью: {_reason(error)}")

    result.text = _merge_list_items(pieces)
    if images:
        result.meta["images"] = images
        result.warnings.append(
            f"изображений в документе: {images} — они не извлекаются, "
            "в тексте оставлены плейсхолдеры"
        )
    try:
        properties = document.core_properties
        core_title = (properties.title or "").strip()
        author = (properties.author or "").strip()
    except Exception:  # pragma: no cover — документ без core.xml
        core_title, author = "", ""
    if author:
        result.meta["author"] = author
    result.title = core_title or heading_title or _first_markdown_heading(result.text) or path.stem
    if result.is_empty:
        result.warnings.append("в DOCX не найдено текста")
    return result


def _merge_list_items(pieces: Sequence[str]) -> str:
    """Соседние пункты списка держим одним абзацем, остальное разделяем пустой строкой."""
    out: List[str] = []
    for piece in pieces:
        if piece.startswith("- ") and out and out[-1].startswith("- "):
            out[-1] = f"{out[-1]}\n{piece}"
            continue
        out.append(piece)
    return "\n\n".join(out)


# -------------------------------------------------------- текст и Markdown ---

#: Однобайтовые кодировки, которые встречаются в старых архивах связи. Порядок
#: здесь ничего не решает — выбор делается по содержимому (см. _score_text):
#: любая из них разбирает ЛЮБЫЕ байты без ошибки, поэтому «первая, которая не
#: упала» всегда давала бы cp1251, а файл в koi8-r молча превращался бы в
#: «оБУФПСЭЙК УФБОДБТФ».
_SINGLE_BYTE = ("cp1251", "koi8-r", "cp866", "iso8859-5", "mac-cyrillic")

#: Знаки, которых в русском техническом тексте не бывает. Их появление —
#: верный признак того, что кодировку выбрали неправильно: типографский блок
#: cp1251 (0x80–0x9F), псевдографика, редкие буквы других славянских языков.
_JUNK = set(
    "\u0402\u0403\u201a\u0453\u201e\u2026\u2020\u2021\u20ac\u2030\u0409\u2039"
    "\u040a\u040c\u040b\u040f\u0452\u2018\u2019\u201c\u201d\u2022\u0459\u045a"
    "\u045c\u045b\u045f\u040e\u045e\u0408\u00a4\u0490\u00a6\u00a7\u00a9\u0404"
    "\u00ab\u00ac\u00ae\u0407\u00b0\u00b1\u0406\u0456\u0491\u00b5\u00b6\u00b7"
    "\u0451\u2116\u0454\u00bb\u0458\u0405\u0455\u0457"
    "\u2591\u2592\u2593\u2502\u2524\u2510\u2514\u2534\u252c\u251c\u2500\u253c"
    "\u2550\u2551\u2557\u255d\u2560\u2563\u256c\u2588\u2584\u258c\u2590\u2580"
)
# «ё» и «№» в русском тексте законны — их из списка мусора убираем.
_JUNK -= set("\u0451\u2116")

#: Ниже этой доли осмысленных знаков считаем, что кодировка не подошла.
_READABLE_MIN = 0.5

#: Самые частые русские двубуквенные сочетания. Это и есть главный признак,
#: по которому кодировка выбирается правильно: koi8-r и cp1251 — перестановки
#: одного и того же алфавита, обе дают «настоящие» русские буквы, и отличить
#: текст от каши по одним буквам нельзя. А вот сочетания «ст», «ов», «ни» в
#: перестановке рассыпаются: в правильном чтении их около 50–75 на сотню букв,
#: в неправильном — единицы.
_BIGRAMS = (
    "ст", "то", "ен", "ни", "ов", "ра", "не", "по", "ко", "на",
    "ре", "ро", "ли", "ва", "ер", "ан", "ор", "та", "ат", "те",
    "ол", "ес", "ка", "ит", "ти", "ых", "ия", "де", "пр", "го",
)
#: Столько сочетаний на букву даёт настоящий русский текст. Всё, что заметно
#: ниже, — признак неверно выбранной кодировки.
_BIGRAM_GOOD = 0.25
#: Меньше этого числа русских букв — судить не по чему (короткая записка,
#: английский документ, таблица из одних чисел).
_BIGRAM_MIN_LETTERS = 20


def _bigram_rate(text: str) -> float | None:
    """Доля частых русских сочетаний. ``None`` — русского текста слишком мало."""
    low = text.lower()
    letters = sum(1 for char in low if "а" <= char <= "я" or char == "ё")
    if letters < _BIGRAM_MIN_LETTERS:
        return None
    return sum(low.count(pair) for pair in _BIGRAMS) * 2 / letters


def _score_text(text: str) -> float:
    """Насколько текст похож на осмысленный. 1.0 — чистый текст, 0.0 — каша.

    Однобайтовые кодировки не дают ошибки декодирования: любой байт во что-то
    да превратится. Отличить настоящий текст от каши можно только по тому, что
    получилось, — этим и занимается функция. Два признака:

    * доля знаков, которых в техническом тексте не бывает (псевдографика,
      типографский блок cp1251, редкие буквы других славянских языков);
    * частота обычных русских буквосочетаний.

    Второй признак и решает дело. Первый ловит cp866 и mac-cyrillic,
    прочитанные как cp1251 (сплошные «Ґ», «§», «€»), но против пары
    koi8-r/cp1251 бессилен: там все буквы настоящие.
    """
    if not text:
        return 0.0
    sample = text[:4000]
    good = junk = 0
    for char in sample:
        if char in _JUNK:
            junk += 1
        elif (char.isspace() or char.isalnum()
              or char in ".,:;!?()[]{}-–—«»\"'/\\+=*%№#@&_|<>$"):
            good += 1
        else:
            junk += 1
    total = good + junk
    if not total:
        return 0.0
    score = good / total

    rate = _bigram_rate(sample)
    if rate is not None:
        score *= max(0.1, min(1.0, rate / _BIGRAM_GOOD))
    return score


def read_text(path: Path) -> Tuple[str, str | None, str | None]:
    """Читает текстовый файл, перебирая кодировки. Возвращает (текст, кодировка, ошибка)."""
    try:
        raw = path.read_bytes()
    except OSError as error:
        return "", None, f"файл не прочитан: {_reason(error)}"
    return decode_bytes(raw)


def decode_bytes(raw: bytes) -> Tuple[str, str | None, str | None]:
    """То же самое для содержимого, уже прочитанного в память.

    Нужно там, где байты берутся не из файла: вложение письма, страница из
    архива сохранённого сайта.
    """
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig", errors="replace"), "utf-8-sig", None
    for bom, encoding in ((b"\xff\xfe", "utf-16-le"), (b"\xfe\xff", "utf-16-be")):
        if raw.startswith(bom):
            return raw[2:].decode(encoding, errors="replace"), encoding, None
    # UTF-8 проверяется строгим разбором: он ошибку даёт, и если файл прошёл —
    # это он и есть.
    try:
        return raw.decode("utf-8"), "utf-8", None
    except UnicodeDecodeError:
        pass

    best_text, best_encoding, best_score = "", "", -1.0
    for encoding in _SINGLE_BYTE:
        try:
            candidate = raw.decode(encoding)
        except UnicodeDecodeError:
            continue
        score = _score_text(candidate)
        if score > best_score:
            best_text, best_encoding, best_score = candidate, encoding, score

    if best_encoding and best_score >= _READABLE_MIN:
        return best_text, best_encoding, None
    if best_encoding:
        return (
            best_text,
            best_encoding,
            f"кодировка определена ненадёжно (выбрана {best_encoding}), "
            "текст может быть искажён — проверьте документ",
        )
    return (
        raw.decode("utf-8", errors="replace"),
        "utf-8/replace",
        "кодировка файла не распознана, часть символов заменена",
    )


def _first_markdown_heading(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip()
    return ""


#: Длиннее этого первая строка на название уже не тянет — это абзац.
_TEXT_TITLE_MAX = 120


def _first_text_line_as_title(text: str) -> str:
    """Первая строка обычного текста, если она похожа на название.

    В .txt разметки нет, и названием документа становилось имя файла:
    паспорт микросхемы назывался «ad9361», методика — «методика_v2_итог».
    Между тем первая строка почти всегда и есть название — надо только
    отличить её от начала текста.
    """
    for line in text.splitlines()[:5]:
        stripped = line.strip()
        if not stripped:
            continue
        if len(stripped) > _TEXT_TITLE_MAX:
            return ""
        # Название не заканчивается точкой и не состоит из одних цифр.
        if stripped.endswith((".", ",", ";", ":")):
            return ""
        if not any(char.isalpha() for char in stripped):
            return ""
        return stripped
    return ""


def _convert_text(path: Path) -> ConvertedDocument:
    """Markdown и обычный текст берём как есть — их уже написал человек."""
    kind = "markdown" if path.suffix.lower() in (".md", ".markdown") else "text"
    result = ConvertedDocument(title=path.stem, meta={"source_format": kind})
    text, encoding, problem = read_text(path)
    if problem:
        result.warnings.append(problem)
    if encoding:
        result.meta["encoding"] = encoding
    result.text = text.replace("\r\n", "\n").replace("\r", "\n")

    from ..corpus import parse_front_matter

    front, body = parse_front_matter(result.text)
    for key, value in front.items():
        result.meta.setdefault(key, value)
    result.title = (front.get("title") or _first_markdown_heading(body)
                    or (_first_text_line_as_title(body) if kind == "text" else "")
                    or path.stem)
    if not result.text.strip():
        result.warnings.append("файл пуст")
    return result


# ---------------------------------------------------------- диспетчер ----

def _register_builtin_converters() -> None:
    """Форматы, которые система разбирает сама, без сторонних инструментов."""
    from . import registry

    registry.register(registry.ConverterSpec(
        name="pdf",
        suffixes=(".pdf",),
        convert=_convert_pdf,
        requires=(registry.Requirement("python", "pymupdf", pip_hint("pymupdf")),),
        note="текстовый слой PDF, заголовки по кеглю, номера страниц",
    ))
    registry.register(registry.ConverterSpec(
        name="docx",
        suffixes=(".docx", ".dotx"),
        convert=_convert_docx,
        requires=(registry.Requirement("python", "docx", pip_hint("python-docx")),),
        note="Word 2007 и новее: заголовки по стилям, таблицы, списки",
    ))
    registry.register(registry.ConverterSpec(
        name="text",
        suffixes=(".md", ".markdown", ".txt"),
        convert=_convert_text,
        note="текст и Markdown, кодировка определяется автоматически",
    ))


_register_builtin_converters()


def supported_suffixes(*, only_available: bool = False) -> Tuple[str, ...]:
    """Все расширения, которые система умеет разбирать сейчас."""
    from . import registry

    return registry.supported_suffixes(only_available=only_available)


def format_support() -> List[Dict[str, Any]]:
    """Состояние поддержки форматов: что доступно, чего не хватает."""
    from . import registry

    return registry.report()


#: Форматы-контейнеры, внутри которых бывают вставленные картинки.
EMBEDDED_IMAGE_SUFFIXES = (
    ".docx", ".dotx", ".pptx", ".potx", ".ppsx", ".xlsx", ".xlsm",
    ".odt", ".ott", ".odp", ".otp", ".ods", ".ots",
)


def embedded_ocr_enabled() -> bool:
    """Распознавать ли картинки внутри документов.

    По умолчанию — да: в технических отчётах половина существенного лежит на
    иллюстрациях (спектрограмма с подписанными частотами, снимок экрана
    анализатора, вклеенная страница методики), и без распознавания всё это
    в базу не попадает вовсе. Выключается переменной окружения, если приём
    большой библиотеки нужно ускорить: REPORTGEN_OCR_EMBEDDED=0.
    """
    return os.environ.get("REPORTGEN_OCR_EMBEDDED", "1").strip().lower() not in (
        "0", "false", "no", "off", ""
    )


def _add_embedded_images(result: "ConvertedDocument", path: Path) -> "ConvertedDocument":
    """Дописать в документ текст, распознанный на вложенных картинках."""
    if path.suffix.lower() not in EMBEDDED_IMAGE_SUFFIXES or not embedded_ocr_enabled():
        return result
    try:
        from .formats.ocr import (  # noqa: PLC0415
            embedded_images_block,
            ocr_embedded_images,
        )

        found, warnings = ocr_embedded_images(path)
    except Exception:  # noqa: BLE001 — распознавание картинок не критично
        return result

    result.warnings.extend(warnings)
    if not found:
        return result
    block = embedded_images_block(found)
    result.text = (result.text.rstrip() + "\n\n" + block) if result.text.strip() else block
    result.meta["embedded_images_read"] = len(found)
    # Прежнее предупреждение «изображения не извлекаются» теперь неверно:
    # они как раз извлечены, и два противоречащих сообщения только путают.
    result.warnings = [
        text for text in result.warnings if "не извлекаются" not in text
    ]
    result.warnings.append(
        f"текст распознан на иллюстрациях: {len(found)} "
        "— числа оттуда сверяйте с оригиналом"
    )
    return result


#: Доля осмысленных знаков, ниже которой текст считается нечитаемым.
MIN_READABLE_SHARE = 0.35

#: Символы-заглушки, которыми разборщики отмечают глифы без соответствия в
#: юникоде: у старых PDF без карты символов ими оказывается весь текст.
_PLACEHOLDER_CHARS = "·•\u00b7\ufffd\u25a0\u25a1\u2591\u2592\u2593?"


def readable_share(text: str) -> float:
    """Какая часть непробельных знаков — буквы или цифры.

    У PDF без карты символов (ToUnicode) текст извлекается, но состоит из
    заглушек: «······ ··········· ·····» вместо «Основы спутниковой связи».
    Формально разбор удался, предупреждений нет, и такой документ молча
    попадает в базу — вместе с заголовком из точек. В библиотеке из старых
    переведённых в PDF книг это встречается регулярно.
    """
    meaningful = [ch for ch in text if not ch.isspace()]
    if not meaningful:
        return 0.0
    good = sum(1 for ch in meaningful if ch.isalnum())
    return good / len(meaningful)


def _repair_result(result: "ConvertedDocument") -> "ConvertedDocument":
    """Починить текст любого формата — и сказать, что именно поправлено.

    Лигатуры, невидимые переносы и латинские буквы внутри русских слов родом
    не из PDF: они так же приходят из DOCX, из старого .doc, из распознанного
    скана и из DjVu. Разборщик PDF правит свою часть сам (пробелы между
    словами и индексы — по геометрии страницы), поэтому повторный проход тут
    ничего у него не отнимает: все починки идемпотентны.
    """
    raw = result.text or ""
    if not raw:
        return result
    fixed = repair_text(raw)
    if fixed == raw:
        return result
    result.text = fixed
    repairs = repair_report(raw, fixed)
    if not repairs:
        return result
    totals = dict(result.meta.get("text_repairs") or {})
    for key, count in repairs.items():
        totals[key] = int(totals.get(key, 0)) + int(count)
    result.meta["text_repairs"] = totals
    note = _repairs_note(repairs)
    if note not in result.warnings:
        result.warnings.append(note)
    return result


def _flag_unreadable(result: "ConvertedDocument", path: Path) -> "ConvertedDocument":
    """Пометить документ, из которого извлеклась бессмыслица.

    Заголовок из заглушек заменяется именем файла: в списке библиотеки строка
    «······ ····» не говорит вообще ничего, а имя файла — говорит.
    """
    text = result.text or ""
    # Короткие документы не проверяем: на десятке знаков доля ничего не значит.
    if len(text.strip()) < 40:
        return result
    share = readable_share(text)
    if share >= MIN_READABLE_SHARE:
        return result
    placeholders = sum(text.count(ch) for ch in _PLACEHOLDER_CHARS)
    result.meta["readable_share"] = round(share, 3)
    if placeholders > len(text) * 0.2:
        result.warnings.append(
            "текст извлёкся заглушками вместо букв: в файле нет карты символов "
            "(ToUnicode). Такой документ ищется только по имени — пересохраните "
            "его или распознайте как скан"
        )
    else:
        result.warnings.append(
            f"текст извлёкся неразборчиво (осмысленных знаков {share:.0%}): "
            "проверьте файл, возможно, нужен скан с распознаванием"
        )
    if not result.title or readable_share(result.title) < MIN_READABLE_SHARE:
        result.title = path.stem
    return result


def convert_file(path: str | Path) -> ConvertedDocument:
    """Конвертирует файл в Markdown по его расширению.

    Функция никогда не бросает исключений на содержимом файла: битый PDF или
    защищённый паролем DOCX дают пустой текст и предупреждение. Это осознанное
    решение — приём каталога на тысячу файлов не должен падать на одном.

    Если формат известен, но инструмент для него не установлен, в
    предупреждении будет сказано, что именно доложить — в изолированном
    контуре это экономит часы выяснений.
    """
    from . import registry

    path = Path(path)
    suffix = path.suffix.lower()
    if not path.is_file():
        result = ConvertedDocument(title=path.stem)
        result.warnings.append(f"файл не найден: {path}")
        return result

    spec = registry.find(suffix)
    if spec is None:
        result = ConvertedDocument(title=path.stem,
                                   meta={"source_format": suffix.lstrip(".")})
        available = ", ".join(registry.supported_suffixes(only_available=True))
        result.warnings.append(
            f"неподдерживаемый формат «{suffix or 'без расширения'}»: "
            f"сейчас доступны {available}"
        )
        return result

    if not spec.is_available():
        result = ConvertedDocument(title=path.stem,
                                   meta={"source_format": suffix.lstrip(".")})
        result.warnings.append(
            f"формат «{suffix}» разбирается конвертером «{spec.name}», но "
            f"{registry.missing_hint(spec)}"
        )
        return result

    try:
        return _flag_unreadable(
            _repair_result(_add_embedded_images(spec.convert(path), path)), path)
    except MissingDependencyError as error:
        result = ConvertedDocument(title=path.stem,
                                   meta={"source_format": suffix.lstrip(".")})
        result.warnings.append(str(error))
        return result
    except Exception as error:  # noqa: BLE001 — один файл не должен ронять приём каталога
        result = ConvertedDocument(title=path.stem,
                                   meta={"source_format": suffix.lstrip(".")})
        result.warnings.append(
            f"конвертер «{spec.name}» не справился с файлом: {_reason(error)}"
        )
        return result


#: Старые имена: на них опираются модули форматов и внешний код.
_read_text = read_text
_clean_line = clean_line


def __getattr__(name: str) -> object:
    """SUPPORTED_SUFFIXES отдаёт настоящий список форматов этой установки.

    Раньше это был кортеж из пяти расширений, зашитый в код. После перехода на
    реестр конвертеров он стал врать: система читает шесть десятков форматов,
    а константа обещала пять. Вычисляем на лету.
    """
    if name == "SUPPORTED_SUFFIXES":
        return supported_suffixes()
    raise AttributeError(f"module {__name__!r} has no attribute {name!r}")
