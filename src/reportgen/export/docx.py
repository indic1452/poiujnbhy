"""Экспорт готового отчёта из Markdown в DOCX по фирменному шаблону.

Слой 7 архитектуры (док. 01, раздел 1.2). К этому моменту текст уже собран
конвейером (:mod:`reportgen.pipeline`) и проверен верификатором
(:mod:`reportgen.verify`) — здесь он только «одевается» в фирменный бланк.
Модуль ничего не считает и ничего не дописывает: источник текста ровно один —
переданный Markdown, поэтому инвариант «ни одного числа мимо факт-пакета»
(док. 01, раздел 1.4) экспортом нарушен быть не может.

Почему собственный конвертер, а не pandoc:

* контур изолирован (док. 07) — рассчитывать на сторонние бинарники нельзя;
* нужен ровно тот поднабор Markdown, который порождает
  :func:`reportgen.pipeline.assemble`: заголовки, таблицы измерений, цитаты
  источников, списки, горизонтальные линии и служебный блок;
* фирменный бланк приходит как ``.docx``-шаблон, а не как набор правил
  оформления: его открывают как основу документа, чтобы подхватились стили,
  поля страницы и колонтитулы отдела.

``python-docx`` импортируется лениво, внутри функций: ядро, CLI и веб-слой
обязаны запускаться там, где пакет не установлен, а понятная ошибка на русском
выдаётся только тому, кто действительно попросил DOCX.

Поддерживаемая разметка::

    # … ######    заголовки шести уровней
    абзац         **жирный**, *курсив*, `моноширинный`, [текст](ссылка)
    | a | b |     таблицы, включая выравнивание (:---, :---:, ---:)
    > цитата      цитаты
    - пункт       маркированные списки (до трёх уровней вложенности)
    1. пункт      нумерованные списки
    ---           горизонтальная линия
    ```           блок кода моноширинным шрифтом
    <!-- … -->    HTML-комментарии пропускаются (в том числе многострочные)

Всё, что не распознано, попадает в документ обычным абзацем: экспорт не имеет
права терять текст отчёта.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Sequence

from ..packages import pip_hint

__all__ = [
    "DRAFT_NOTICE",
    "ExportOptions",
    "MissingDependencyError",
    "export_report",
    "footer_for",
    "is_draft",
    "markdown_to_docx",
]

# --------------------------------------------------------------- параметры ---

#: Надпись в верхнем колонтитуле неутверждённого отчёта (док. 01, инвариант 4).
DRAFT_NOTICE = "ЧЕРНОВИК — требует проверки инженера"

#: Текст, который Word покажет вместо оглавления, пока поле не обновлено.
TOC_PLACEHOLDER = "Оглавление будет построено Word при открытии документа."
TOC_HINT = "Если оглавление не обновилось автоматически — выделите его и нажмите F9."
TOC_TITLE = "Содержание"
TOC_INSTRUCTION = r' TOC \o "1-3" \h \z \u '
PAGE_INSTRUCTION = " PAGE "

#: Умолчания оформления, когда фирменный шаблон не задан.
BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12
MONO_FONT = "Courier New"
MONO_SIZE_PT = 10
MARGIN_CM = 2.0
LINE_SPACING = 1.15
SPACE_AFTER_PT = 6

#: Статусы отчёта, при которых документ уже не черновик (см. store/schema.sql).
APPROVED_STATUSES = frozenset({"approved", "утверждён", "утвержден", "final"})

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_RULE_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")
_FENCE_RE = re.compile(r"^\s*(?:`{3,}|~{3,})\s*\S*\s*$")
_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_ORDERED_RE = re.compile(r"^(\s*)\d{1,3}[.)]\s+(.*)$")
_QUOTE_RE = re.compile(r"^\s*>\s?(.*)$")
_TABLE_LINE_RE = re.compile(r"^\s*\|.*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*:?-*:?\s*\|?\s*$")
_HARD_BREAK_RE = re.compile(r"(?:\s{2,}|\\)$")
_NUMBERING_RE = re.compile(r"^\s*(?:\d+[.)]\s*)+")
_APPENDIX_RE = re.compile(r"^(?:\d+[.)]\s*)*приложени[ея]\b", re.IGNORECASE)
_TOC_TITLES = frozenset({"содержание", "оглавление", "contents", "table of contents"})

#: Разбор начертаний внутри абзаца. Подчёркивания требуют границы слова, иначе
#: пострадают технические идентификаторы вида ``facts_digest`` и ``ЗАКАЗЧИК_07``.
_INLINE_RE = re.compile(
    # Экранированный знак идёт первым: «\*» — это звёздочка, а не начало
    # курсива. Без этого номер группы «*1274*» и модель «Р_168_М» из
    # факт-пакета приезжали в документ без части знаков.
    r"(?P<esc>\\[\\`*_\[\]])"
    r"|(?P<mono>`[^`\n]+`)"
    r"|(?P<link>\[(?P<link_text>[^\]\n]+)\]\((?P<link_url>[^)\s]+)\))"
    r"|(?P<strong_em>\*\*\*[^*\n]+\*\*\*)"
    r"|(?P<strong>\*\*[^*\n]+\*\*)"
    r"|(?P<strong_u>(?<!\w)__[^_\n]+__(?!\w))"
    r"|(?P<em>\*[^*\n]+\*)"
    r"|(?P<em_u>(?<!\w)_[^_\n]+_(?!\w))"
)


class MissingDependencyError(ImportError):
    """Не установлен пакет, без которого экспорт невозможен.

    Наследуется от :class:`ImportError`, потому что веб-слой
    (``web/api.py``) отличает «функциональность недоступна» от «сборка
    документа не удалась» именно по типу исключения.
    """


@dataclass
class ExportOptions:
    """Параметры выгрузки одного документа.

    :param template: фирменный бланк ``.docx``; его стили, поля и колонтитулы
        становятся основой документа. ``None`` — собрать документ с нуля.
    :param draft: документ ещё не подписан инженером; в верхний колонтитул
        добавляется :data:`DRAFT_NOTICE`.
    :param footer_text: строка нижнего колонтитула (рядом с номером страницы).
    :param page_break_before_appendix: начинать приложения с новой страницы.
    :param toc: вставить поле оглавления, которое Word обновит при открытии.
    """

    template: Path | None = None
    draft: bool = True
    footer_text: str = ""
    page_break_before_appendix: bool = True
    toc: bool = True

    def __post_init__(self) -> None:
        if self.template is not None and not isinstance(self.template, Path):
            self.template = Path(self.template)


# ------------------------------------------------------------ разбор текста ---

@dataclass(frozen=True)
class _Span:
    """Кусок абзаца с уже разобранным начертанием."""

    text: str
    bold: bool = False
    italic: bool = False
    mono: bool = False


@dataclass
class _ListItem:
    """Пункт списка: отступ в исходной разметке, вид маркера, текст."""

    indent: int
    ordered: bool
    text: str


@dataclass
class _Block:
    """Разобранный блок Markdown — единица, которая станет абзацем или таблицей.

    Поля заполняются в зависимости от ``kind``: ``heading`` использует
    ``level``/``text``, ``paragraph``/``quote``/``code`` — ``lines``,
    ``table`` — ``rows``/``aligns``, ``list`` — ``items``, ``rule`` и ``toc`` —
    ничего сверх ``kind`` (у ``toc`` дополнительно ``level``/``text``
    заголовка раздела).
    """

    kind: str
    level: int = 0
    text: str = ""
    lines: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    aligns: List[str | None] = field(default_factory=list)
    items: List[_ListItem] = field(default_factory=list)


def _strip_html_comments(text: str) -> str:
    """Убирает HTML-комментарии, не трогая содержимое блоков кода.

    В отчёте комментарии несут служебную информацию (маркеры страниц из
    слоя приёма, пояснение к служебному блоку) — в готовый документ они
    попадать не должны.
    """
    result: List[str] = []
    in_fence = False
    in_comment = False
    for raw in text.splitlines():
        if not in_comment and _FENCE_RE.match(raw):
            in_fence = not in_fence
            result.append(raw)
            continue
        if in_fence:
            result.append(raw)
            continue

        line = raw
        if in_comment:
            end = line.find("-->")
            if end == -1:
                continue
            line = line[end + 3:]
            in_comment = False

        while "<!--" in line:
            start = line.find("<!--")
            end = line.find("-->", start + 4)
            if end == -1:
                line = line[:start]
                in_comment = True
                break
            line = line[:start] + line[end + 3:]

        if line.strip() or not raw.strip():
            result.append(line)
    return "\n".join(result)


def _split_cells(line: str) -> List[str]:
    """Ячейки строки таблицы. Экранированный разделитель ``\\|`` остаётся текстом."""
    body = line.strip()
    if body.startswith("|"):
        body = body[1:]
    if body.endswith("|") and not body.endswith("\\|"):
        body = body[:-1]
    cells = re.split(r"(?<!\\)\|", body)
    return [cell.replace("\\|", "|").strip() for cell in cells]


def _parse_aligns(line: str) -> List[str | None]:
    """Выравнивание столбцов из строки-разделителя таблицы."""
    aligns: List[str | None] = []
    for cell in _split_cells(line):
        left = cell.startswith(":")
        right = cell.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        elif left:
            aligns.append("left")
        else:
            aligns.append(None)
    return aligns


def _parse_blocks(markdown: str) -> List[_Block]:
    """Разбирает Markdown на блоки. Неизвестная разметка становится абзацем."""
    lines = _strip_html_comments(markdown).splitlines()
    blocks: List[_Block] = []
    index = 0
    total = len(lines)

    while index < total:
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if _FENCE_RE.match(line):
            index += 1
            code: List[str] = []
            while index < total and not _FENCE_RE.match(lines[index]):
                code.append(lines[index])
                index += 1
            index += 1  # закрывающая ограда; её отсутствие ошибкой не считаем
            blocks.append(_Block(kind="code", lines=code))
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            blocks.append(_Block(
                kind="heading",
                level=len(heading.group(1)),
                text=heading.group(2).strip(),
            ))
            index += 1
            continue

        if _RULE_RE.match(line):
            blocks.append(_Block(kind="rule"))
            index += 1
            continue

        if (
            _TABLE_LINE_RE.match(line)
            and index + 1 < total
            and _TABLE_SEP_RE.match(lines[index + 1])
        ):
            aligns = _parse_aligns(lines[index + 1])
            rows = [_split_cells(line)]
            index += 2
            while index < total and _TABLE_LINE_RE.match(lines[index]):
                rows.append(_split_cells(lines[index]))
                index += 1
            blocks.append(_Block(kind="table", rows=rows, aligns=aligns))
            continue

        if _QUOTE_RE.match(line):
            quoted: List[str] = []
            while index < total and _QUOTE_RE.match(lines[index]):
                quoted.append(_QUOTE_RE.match(lines[index]).group(1))
                index += 1
            blocks.append(_Block(kind="quote", lines=quoted))
            continue

        if _BULLET_RE.match(line) or _ORDERED_RE.match(line):
            items: List[_ListItem] = []
            while index < total and lines[index].strip():
                bullet = _BULLET_RE.match(lines[index])
                ordered = _ORDERED_RE.match(lines[index])
                if bullet:
                    items.append(_ListItem(len(bullet.group(1)), False, bullet.group(2).strip()))
                elif ordered:
                    items.append(_ListItem(len(ordered.group(1)), True, ordered.group(2).strip()))
                elif items:
                    # Продолжение пункта на следующей строке.
                    items[-1].text = f"{items[-1].text} {lines[index].strip()}".strip()
                else:
                    break
                index += 1
            blocks.append(_Block(kind="list", items=items))
            continue

        paragraph: List[str] = []
        while index < total and lines[index].strip():
            current = lines[index]
            if (
                _HEADING_RE.match(current)
                or _RULE_RE.match(current)
                or _QUOTE_RE.match(current)
                or _FENCE_RE.match(current)
                or _TABLE_LINE_RE.match(current)
            ):
                break
            paragraph.append(current)
            index += 1
        if paragraph:
            blocks.append(_Block(kind="paragraph", lines=paragraph))
        else:
            # Строка распознана как начало другого блока — вернёмся к ней снаружи.
            blocks.append(_Block(kind="paragraph", lines=[lines[index]]))
            index += 1

    return blocks


def _logical_lines(lines: Sequence[str]) -> List[str]:
    """Склеивает строки абзаца, уважая жёсткий перенос (два пробела в конце)."""
    result: List[str] = []
    current: List[str] = []
    for line in lines:
        text = line.strip()
        if not text:
            continue
        current.append(text)
        if _HARD_BREAK_RE.search(line.rstrip("\n")):
            result.append(" ".join(current))
            current = []
    if current:
        result.append(" ".join(current))
    return result


def _split_inline(text: str) -> List[_Span]:
    """Разбирает **жирный**, *курсив*, `моноширинный` и ссылки внутри абзаца."""
    spans: List[_Span] = []
    position = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            spans.append(_Span(text[position:match.start()]))
        groups = match.groupdict()
        if groups["esc"] is not None:
            # «\*» — это звёздочка. Сама косая черта в документ не идёт.
            spans.append(_Span(groups["esc"][1]))
        elif groups["mono"] is not None:
            spans.append(_Span(groups["mono"][1:-1], mono=True))
        elif groups["link"] is not None:
            label = groups["link_text"]
            url = groups["link_url"]
            spans.append(_Span(label if url.startswith("#") else f"{label} ({url})"))
        elif groups["strong_em"] is not None:
            spans.append(_Span(groups["strong_em"][3:-3], bold=True, italic=True))
        elif groups["strong"] is not None:
            spans.append(_Span(groups["strong"][2:-2], bold=True))
        elif groups["strong_u"] is not None:
            spans.append(_Span(groups["strong_u"][2:-2], bold=True))
        elif groups["em"] is not None:
            spans.append(_Span(groups["em"][1:-1], italic=True))
        elif groups["em_u"] is not None:
            spans.append(_Span(groups["em_u"][1:-1], italic=True))
        position = match.end()
    if position < len(text):
        spans.append(_Span(text[position:]))
    return [span for span in spans if span.text]


def _normalize_title(text: str) -> str:
    """Заголовок без нумерации и регистра — для сравнения со служебными именами."""
    return _NUMBERING_RE.sub("", text).strip().strip(".").lower()


def _is_appendix(text: str) -> bool:
    return bool(_APPENDIX_RE.match(text.strip()))


def _with_toc(blocks: Sequence[_Block]) -> List[_Block]:
    """Ставит поле оглавления после титульного блока.

    Если конвейер уже вывел раздел «Содержание» со списком разделов, он
    заменяется полем: два оглавления подряд в документе не нужны.
    """
    blocks = list(blocks)
    for index, block in enumerate(blocks):
        if block.kind == "heading" and _normalize_title(block.text) in _TOC_TITLES:
            tail = index + 1
            if tail < len(blocks) and blocks[tail].kind == "list":
                tail += 1
            toc = _Block(kind="toc", level=block.level, text=block.text)
            return blocks[:index] + [toc] + blocks[tail:]

    for index, block in enumerate(blocks):
        if block.kind == "heading" and block.level >= 2:
            toc = _Block(kind="toc", level=2, text=TOC_TITLE)
            return blocks[:index] + [toc] + blocks[index:]

    return blocks + [_Block(kind="toc", level=1, text=TOC_TITLE)]


# --------------------------------------------------------- доступ к docx ---

@dataclass(frozen=True)
class _DocxApi:
    """Кусочки python-docx, нужные экспорту. Собираются одним ленивым импортом."""

    Document: Any
    Pt: Any
    Cm: Any
    align: Any
    element: Any
    qn: Any


def _docx_api() -> _DocxApi:
    """Ленивая загрузка python-docx с понятной ошибкой в изолированном контуре."""
    try:
        from docx import Document  # noqa: PLC0415
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
        from docx.oxml import OxmlElement  # noqa: PLC0415
        from docx.oxml.ns import qn  # noqa: PLC0415
        from docx.shared import Cm, Pt  # noqa: PLC0415
    except ImportError as error:
        raise MissingDependencyError(
            "экспорт в DOCX недоступен: не установлен пакет python-docx. "
            "Отчёт можно забрать в виде Markdown — эта кнопка работает без "
            "него. Чтобы вернуть выгрузку в DOCX, администратору системы "
            "нужно выполнить на сервере: " + pip_hint("python-docx")
        ) from error
    return _DocxApi(
        Document=Document,
        Pt=Pt,
        Cm=Cm,
        align=WD_ALIGN_PARAGRAPH,
        element=OxmlElement,
        qn=qn,
    )


def _open_document(api: _DocxApi, options: ExportOptions) -> Any:
    """Открывает фирменный бланк или создаёт документ с умолчаниями отдела."""
    if options.template is None:
        document = api.Document()
        _apply_defaults(api, document)
        return document

    template = Path(options.template)
    if not template.is_file():
        raise FileNotFoundError(
            f"шаблон DOCX не найден: {template}. Проверьте настройку "
            f"REPORTGEN_DOCX_TEMPLATE или уберите её, чтобы собрать документ "
            f"без фирменного бланка."
        )
    document = api.Document(str(template))
    _drop_trailing_empty_paragraphs(document)
    return document


def _apply_defaults(api: _DocxApi, document: Any) -> None:
    """Умолчания для документа без бланка: Times New Roman 12, поля 2 см, 1.15."""
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = api.Pt(BODY_SIZE_PT)
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(api.qn(attribute), BODY_FONT)

    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING
    paragraph_format.space_after = api.Pt(SPACE_AFTER_PT)

    for section in document.sections:
        section.top_margin = api.Cm(MARGIN_CM)
        section.bottom_margin = api.Cm(MARGIN_CM)
        section.left_margin = api.Cm(MARGIN_CM)
        section.right_margin = api.Cm(MARGIN_CM)


_SECT_PR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/" \
           "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"


def _drop_trailing_empty_paragraphs(document: Any) -> None:
    """Убирает пустые абзацы в хвосте бланка, чтобы отчёт не начинался с пропусков.

    Абзац, несущий описание секции (``w:sectPr``), не трогаем: вместе с ним
    из документа исчезли бы поля страницы и ориентация листа.
    """
    while document.paragraphs:
        paragraph = document.paragraphs[-1]
        if paragraph.text.strip() or paragraph._p.find(_SECT_PR) is not None:
            break
        if paragraph._p.getparent() is None:
            break
        paragraph._p.getparent().remove(paragraph._p)


def _add_field(api: _DocxApi, paragraph: Any, instruction: str, placeholder: str) -> None:
    """Вставляет поле Word (``TOC``, ``PAGE``) с текстом-заглушкой.

    Поле собирается в развёрнутой форме (begin / instrText / separate /
    результат / end): так заглушка остаётся обычным текстом абзаца и видна
    даже в тех просмотрщиках, которые полей не обновляют.
    """
    begin = paragraph.add_run()
    marker = api.element("w:fldChar")
    marker.set(api.qn("w:fldCharType"), "begin")
    begin._r.append(marker)

    instruction_run = paragraph.add_run()
    instruction_element = api.element("w:instrText")
    instruction_element.set(api.qn("xml:space"), "preserve")
    instruction_element.text = instruction
    instruction_run._r.append(instruction_element)

    separate = paragraph.add_run()
    separator = api.element("w:fldChar")
    separator.set(api.qn("w:fldCharType"), "separate")
    separate._r.append(separator)

    if placeholder:
        paragraph.add_run(placeholder)

    end = paragraph.add_run()
    closing = api.element("w:fldChar")
    closing.set(api.qn("w:fldCharType"), "end")
    end._r.append(closing)


def _request_field_update(api: _DocxApi, document: Any) -> None:
    """Просит Word обновить поля при открытии — иначе оглавление останется пустым."""
    settings = document.settings.element
    if settings.find(api.qn("w:updateFields")) is not None:
        return
    update = api.element("w:updateFields")
    update.set(api.qn("w:val"), "true")
    compat = settings.find(api.qn("w:compat"))
    if compat is not None:
        compat.addprevious(update)
    else:
        settings.append(update)


# ------------------------------------------------------------ сборка тела ---

class _Writer:
    """Переносит разобранные блоки Markdown в документ Word.

    Стили берутся из шаблона по имени; если фирменный бланк какой-то стиль не
    объявляет, оформление воспроизводится вручную (жирный заголовок, отступ
    цитаты, маркер списка). Экспорт не должен падать из-за чужого бланка.
    """

    def __init__(self, api: _DocxApi, document: Any, options: ExportOptions) -> None:
        self.api = api
        self.document = document
        self.options = options
        self._styles: Dict[str, bool] = {}
        self._written = 0

    # -- служебное ----------------------------------------------------------

    def _has_style(self, name: str) -> bool:
        if name not in self._styles:
            try:
                self.document.styles[name]
                self._styles[name] = True
            except KeyError:
                self._styles[name] = False
        return self._styles[name]

    def _style(self, *names: str) -> str | None:
        for name in names:
            if self._has_style(name):
                return name
        return None

    def _paragraph(self, *names: str) -> Any:
        return self.document.add_paragraph(style=self._style(*names))

    def _runs(self, paragraph: Any, text: str, *, bold: bool = False) -> None:
        """Добавляет в абзац текст с разобранным начертанием.

        ``bold`` включает жирность всей строки (шапка таблицы). Значение
        ``False`` в runs не выставляется: иначе оно перебило бы жирность,
        заданную стилем заголовка в фирменном бланке.
        """
        for span in _split_inline(text):
            run = paragraph.add_run(span.text)
            if span.bold or bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.mono:
                run.font.name = MONO_FONT
                run.font.size = self.api.Pt(MONO_SIZE_PT)

    def _multiline(self, paragraph: Any, lines: Sequence[str]) -> None:
        """Логические строки одного абзаца, разделённые мягким переносом."""
        for number, line in enumerate(_logical_lines(lines)):
            if number:
                paragraph.add_run().add_break()
            self._runs(paragraph, line)

    # -- блоки --------------------------------------------------------------

    def write(self, blocks: Sequence[_Block]) -> None:
        handlers = {
            "heading": self._write_heading,
            "paragraph": self._write_paragraph,
            "table": self._write_table,
            "quote": self._write_quote,
            "list": self._write_list,
            "rule": self._write_rule,
            "code": self._write_code,
            "toc": self._write_toc,
        }
        for block in blocks:
            handlers.get(block.kind, self._write_paragraph)(block)
            self._written += 1

    def _write_heading(self, block: _Block) -> None:
        level = max(1, min(block.level, 6))
        style = self._style(f"Heading {level}")
        paragraph = self.document.add_paragraph(style=style)
        if (
            self.options.page_break_before_appendix
            and self._written
            and _is_appendix(block.text)
        ):
            paragraph.paragraph_format.page_break_before = True
        self._runs(paragraph, block.text, bold=style is None)
        if style is None:
            paragraph.paragraph_format.space_before = self.api.Pt(12)
            for run in paragraph.runs:
                run.font.size = self.api.Pt(max(BODY_SIZE_PT, 18 - 2 * level))

    def _write_paragraph(self, block: _Block) -> None:
        lines = block.lines or ([block.text] if block.text else [])
        if not _logical_lines(lines):
            return
        self._multiline(self.document.add_paragraph(), lines)

    def _write_quote(self, block: _Block) -> None:
        if not _logical_lines(block.lines):
            return
        style = self._style("Quote", "Intense Quote")
        paragraph = self.document.add_paragraph(style=style)
        self._multiline(paragraph, block.lines)
        if style is None:
            paragraph.paragraph_format.left_indent = self.api.Cm(1)
            for run in paragraph.runs:
                run.italic = True

    def _write_list(self, block: _Block) -> None:
        depths = sorted({item.indent for item in block.items})
        ranks = {indent: min(number, 2) for number, indent in enumerate(depths)}
        counters: Dict[int, int] = {}
        for item in block.items:
            if not item.text:
                continue
            depth = ranks.get(item.indent, 0)
            base = "List Number" if item.ordered else "List Bullet"
            suffix = "" if depth == 0 else f" {depth + 1}"
            style = self._style(f"{base}{suffix}", base)
            paragraph = self.document.add_paragraph(style=style)
            if style is None:
                # Бланк без списочных стилей: рисуем маркер и отступ руками.
                counters[depth] = counters.get(depth, 0) + 1
                marker = f"{counters[depth]}. " if item.ordered else "• "
                paragraph.add_run(marker)
                paragraph.paragraph_format.left_indent = self.api.Cm(0.75 * (depth + 1))
            self._runs(paragraph, item.text)

    def _write_rule(self, _block: _Block) -> None:
        paragraph = self.document.add_paragraph()
        properties = paragraph._p.get_or_add_pPr()
        borders = self.api.element("w:pBdr")
        bottom = self.api.element("w:bottom")
        bottom.set(self.api.qn("w:val"), "single")
        bottom.set(self.api.qn("w:sz"), "6")
        bottom.set(self.api.qn("w:space"), "1")
        bottom.set(self.api.qn("w:color"), "auto")
        borders.append(bottom)
        properties.append(borders)

    def _write_code(self, block: _Block) -> None:
        style = self._style("No Spacing")
        for line in block.lines or [""]:
            paragraph = self.document.add_paragraph(style=style)
            run = paragraph.add_run(line)
            run.font.name = MONO_FONT
            run.font.size = self.api.Pt(MONO_SIZE_PT)
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(self.api.qn("w:cs"), MONO_FONT)
            paragraph.paragraph_format.space_after = self.api.Pt(0)

    def _write_toc(self, block: _Block) -> None:
        if block.text:
            level = max(1, min(block.level or 2, 6))
            heading = self.document.add_paragraph(style=self._style(f"Heading {level}"))
            self._runs(heading, block.text)
        field_paragraph = self.document.add_paragraph()
        _add_field(self.api, field_paragraph, TOC_INSTRUCTION, TOC_PLACEHOLDER)
        hint = self.document.add_paragraph()
        run = hint.add_run(TOC_HINT)
        run.italic = True
        run.font.size = self.api.Pt(BODY_SIZE_PT - 2)

    def _write_table(self, block: _Block) -> None:
        """Таблица измерений. Пустые строки и рваные строки не должны ронять экспорт."""
        rows = [row for row in block.rows if any(cell.strip() for cell in row)]
        width = max((len(row) for row in rows), default=0)
        if not rows or width == 0:
            return

        table = self.document.add_table(rows=len(rows), cols=width)
        style = self._style("Table Grid")
        if style is not None:
            table.style = style
        table.autofit = True

        for row_number, row in enumerate(rows):
            for column in range(width):
                cell_text = row[column] if column < len(row) else ""
                paragraph = table.cell(row_number, column).paragraphs[0]
                alignment = self._alignment(block.aligns, column)
                if alignment is not None:
                    paragraph.alignment = alignment
                self._runs(paragraph, cell_text, bold=row_number == 0)

        if len(rows) > 1:
            self._repeat_header(table.rows[0])

    def _alignment(self, aligns: Sequence[str | None], column: int) -> Any:
        if column >= len(aligns) or aligns[column] is None:
            return None
        return {
            "left": self.api.align.LEFT,
            "center": self.api.align.CENTER,
            "right": self.api.align.RIGHT,
        }[aligns[column]]

    def _repeat_header(self, row: Any) -> None:
        """Шапка таблицы повторяется на каждой странице — таблицы измерений длинные."""
        properties = row._tr.get_or_add_trPr()
        if properties.find(self.api.qn("w:tblHeader")) is None:
            properties.append(self.api.element("w:tblHeader"))


# ---------------------------------------------------------- колонтитулы ---

def _writable_paragraph(container: Any) -> Any:
    """Абзац колонтитула, куда можно писать, не затирая фирменный бланк."""
    container.is_linked_to_previous = False
    paragraphs = container.paragraphs
    if paragraphs and not paragraphs[0].text.strip() and not paragraphs[0].runs:
        return paragraphs[0]
    return container.add_paragraph()


def _apply_running_titles(api: _DocxApi, document: Any, options: ExportOptions) -> None:
    """Верхний колонтитул с отметкой черновика, нижний — подпись и номер страницы."""
    for section in document.sections:
        if options.draft:
            paragraph = _writable_paragraph(section.header)
            paragraph.alignment = api.align.CENTER
            run = paragraph.add_run(DRAFT_NOTICE)
            run.bold = True
            run.font.size = api.Pt(BODY_SIZE_PT - 2)

        footer = _writable_paragraph(section.footer)
        footer.alignment = api.align.CENTER
        if options.footer_text:
            run = footer.add_run(options.footer_text)
            run.font.size = api.Pt(BODY_SIZE_PT - 2)
            footer.add_run("   ·   ").font.size = api.Pt(BODY_SIZE_PT - 2)
        page = footer.add_run("с. ")
        page.font.size = api.Pt(BODY_SIZE_PT - 2)
        _add_field(api, footer, PAGE_INSTRUCTION, "1")


# ------------------------------------------------------------ публичный API ---

def markdown_to_docx(
    markdown: str,
    path: str | Path,
    options: ExportOptions | None = None,
) -> Path:
    """Собирает DOCX по готовому Markdown отчёта и возвращает путь к файлу.

    :param markdown: текст отчёта (``ReportResult.markdown`` или правка инженера);
    :param path: куда сохранить документ; каталог создаётся при необходимости;
    :param options: параметры выгрузки; ``None`` — :class:`ExportOptions` по умолчанию.

    Требует установленного ``python-docx``: при его отсутствии поднимается
    :class:`MissingDependencyError` с подсказкой на русском.
    """
    options = options or ExportOptions()
    api = _docx_api()
    document = _open_document(api, options)

    blocks = _parse_blocks(markdown)
    if options.toc:
        blocks = _with_toc(blocks)

    _Writer(api, document, options).write(blocks)
    _apply_running_titles(api, document, options)
    if options.toc:
        _request_field_update(api, document)

    target = Path(path)
    if target.parent and not target.parent.exists():
        target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    return target


def is_draft(status: str) -> bool:
    """Считается ли отчёт с таким статусом черновиком (см. store/schema.sql)."""
    return str(status).strip().lower() not in APPROVED_STATUSES


def footer_for(case_id: str, incoming_no: str = "", outgoing_no: str = "",
               note: str = "") -> str:
    """Подпись в нижнем колонтитуле документа.

    Документ уходит адресату, а не остаётся в системе. В колонтитуле должны
    стоять те номера, по которым его найдут в делопроизводстве: входящий, на
    который отвечаем, и исходящий, под которым ответ ушёл. Учётный номер
    системы — запасной вариант: снаружи он никому ничего не говорит.
    """
    parts = []
    outgoing_no = (outgoing_no or "").strip()
    incoming_no = (incoming_no or "").strip()
    case_id = (case_id or "").strip()
    if outgoing_no:
        parts.append(f"исх. {outgoing_no}")
    if incoming_no:
        parts.append(f"на вх. {incoming_no}")
    if not parts and case_id:
        parts.append(f"по обращению {case_id}")
    строка = "Технический отчёт " + " ".join(parts) if parts else "Технический отчёт"
    # Приписка отдела из настроек (report_footer). Настройка была в образце
    # настроек и в документации, но не читалась никем: отдел вписывал туда
    # свою строку и не находил её в выгруженном отчёте.
    note = (note or "").strip()
    if note:
        строка = f"{строка} · {note}"
    return строка


def export_report(
    markdown: str,
    path: str | Path,
    *,
    case_id: str = "",
    incoming_no: str = "",
    outgoing_no: str = "",
    status: str = "draft",
    template: str | Path | None = None,
    note: str = "",
) -> Path:
    """Выгрузка отчёта одной строкой — обёртка для веб-слоя и CLI.

    Сама заполняет :class:`ExportOptions`: отметку черновика по ``status``
    (черновик — всё, что не утверждено) и подпись колонтитула по номерам
    письма — входящему и исходящему.
    """
    options = ExportOptions(
        template=Path(template) if template else None,
        draft=is_draft(status),
        footer_text=footer_for(case_id, incoming_no, outgoing_no, note),
    )
    return markdown_to_docx(markdown, path, options)
