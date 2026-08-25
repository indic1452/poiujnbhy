"""Тесты слоя приёма документов библиотеки.

Файлы для проверки собираются программно: DOCX — через python-docx, PDF —
через PyMuPDF (TextWriter, потому что базовые шрифты PDF не умеют кириллицу).
Ни один тест не ходит в сеть и не требует внешних сервисов.
"""

import hashlib
import shutil
import tempfile
import textwrap
import unittest
from pathlib import Path
from unittest import mock

import _bootstrap  # noqa: F401

from reportgen import corpus
from reportgen.ingest import convert as convert_module
from reportgen.ingest import (
    IngestResult,
    chunks_from_markdown,
    convert_file,
    guess_doc_type,
    ingest_directory,
    ingest_path,
    remove_document,
    sha256_file,
)
from reportgen.store import Database
from reportgen.store.repo import Repositories

try:  # PyMuPDF нужен только тестам разбора PDF
    import pymupdf
except ImportError:  # pragma: no cover — окружение без пакета
    pymupdf = None

try:
    import docx
except ImportError:  # pragma: no cover — окружение без пакета
    docx = None

EXAMPLES = Path(__file__).resolve().parents[1] / "examples" / "corpus"

LONG_TEXT = (
    "Измерения выполнены анализатором спектра в полосе пропускания 30 кГц "
    "при усреднении по 64 сегментам и подтверждены повторным прогоном. "
    "Условия измерений зафиксированы в протоколе и приложены к отчёту. "
    "Отклонений от методики контроля излучения передатчика не выявлено."
)


# ---------------------------------------------------------- вспомогательное ---

def make_pdf(path, pages):
    """Собирает PDF: pages — список страниц, страница — список пар (строка, кегль).

    Длинные строки переносятся вручную: в PDF нет переноса по ширине, текст за
    границей страницы просто теряется. Строки одного абзаца ставятся с
    межстрочным интервалом, по которому PyMuPDF потом соберёт их в один блок.
    """
    document = pymupdf.open()
    font = pymupdf.Font("helv")
    for lines in pages:
        page = document.new_page()
        writer = pymupdf.TextWriter(page.rect)
        offset = 70.0
        for text, size in lines:
            wrapped = textwrap.wrap(text, width=64) or [text]
            for line in wrapped:
                writer.append((60, offset), line, font=font, fontsize=size)
                offset += size * 1.2
            offset += size
        if lines:
            writer.write_text(page)
    document.save(str(path))
    document.close()


def make_docx(path):
    """Собирает DOCX с заголовками, таблицей, списками и подписью к рисунку."""
    document = docx.Document()
    document.core_properties.title = "Отчёт по обращению 2024-118"
    document.add_heading("Отчёт по обращению 2024-118", level=1)
    document.add_paragraph("Заказчик сообщил о срывах связи в утренние часы. " + LONG_TEXT)
    document.add_heading("Результаты измерений", level=2)
    table = document.add_table(rows=3, cols=3)
    rows = [
        ["Параметр", "Значение", "Единица"],
        ["SNR", "13.7", "дБ"],
        ["EVM | допуск", "4.2", "%"],
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    document.add_paragraph("Рис. 1 — спектр принимаемого сигнала", style="Caption")
    document.add_paragraph("проверить полосовой фильтр", style="List Bullet")
    document.add_paragraph("заменить кабельную сборку", style="List Number")
    document.save(str(path))


class TempCase(unittest.TestCase):
    """Общий временный каталог и база в памяти."""

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp(prefix="reportgen-ingest-"))
        self.addCleanup(shutil.rmtree, self.tmp, ignore_errors=True)
        self.db = Database(":memory:")
        self.addCleanup(self.db.close)
        self.repos = Repositories(self.db)

    def write(self, relative, text, encoding="utf-8"):
        path = self.tmp / relative
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding=encoding)
        return path


# ------------------------------------------------------------- конвертация ---

class ConvertTextTests(TempCase):
    def test_markdown_kept_as_is(self):
        source = "---\ntitle: Методика ОБВ\nyear: 2020\n---\n# Методика ОБВ\n\nТело.\n"
        path = self.write("standards/obw.md", source)
        converted = convert_file(path)
        self.assertEqual(converted.text, source)
        self.assertEqual(converted.title, "Методика ОБВ")
        self.assertEqual(converted.meta["source_format"], "markdown")
        self.assertEqual(converted.meta["year"], "2020")
        self.assertFalse(converted.warnings)

    def test_plain_text_in_cp1251(self):
        path = self.tmp / "regulations" / "заметка.txt"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes("Регламент оформления отчётов.".encode("cp1251"))
        converted = convert_file(path)
        self.assertIn("Регламент оформления", converted.text)
        self.assertEqual(converted.meta["encoding"], "cp1251")
        self.assertEqual(converted.meta["source_format"], "text")

    def test_title_comes_from_the_first_line_of_plain_text(self):
        """В .txt разметки нет, но название почти всегда в первой строке.

        Иначе паспорт микросхемы называется «ad9361», а методика —
        «методика_v2_итог»: по такому названию документ не найти, и тип по
        нему не определить (название весит втрое против текста).
        """
        path = self.write("datasheets/ad9361.txt",
                          "AD9361 RF Agile Transceiver Data Sheet\n\nFEATURES\n"
                          "RF 2x2 transceiver with integrated DAC and ADC.\n")
        self.assertEqual("AD9361 RF Agile Transceiver Data Sheet",
                         convert_file(path).title)

    def test_first_line_that_is_not_a_title_is_not_used(self):
        for name, body in (
            ("абзац.txt", "Очень длинное начало документа, которое на самом деле "
                          "является первым абзацем текста, а вовсе не названием, "
                          "и потому названием становиться не должно\n"),
            ("фраза.txt", "Позвонить в понедельник, уточнить срок поставки.\n"),
            ("номер.txt", "2024-017\n\nОтчёт по обращению\n"),
        ):
            with self.subTest(name=name):
                path = self.write(f"literature/{name}", body)
                self.assertEqual(Path(name).stem, convert_file(path).title)

    def test_legacy_cyrillic_encodings_are_recognised(self):
        """Старые архивы приносят не только cp1251.

        Однобайтовые кодировки ошибку декодирования НЕ дают: любой байт во
        что-то да превратится. Поэтому «первая, которая не упала» всегда
        давала бы cp1251, и файл в koi8-r молча становился бы «оБУФПСЭЙК
        УФБОДБТФ» — с виду русский текст, в поиске бесполезный.
        """
        source = ("Настоящий стандарт распространяется на цифровые "
                  "радиорелейные линии связи и устанавливает требования "
                  "к показателям качества")
        for encoding in ("cp1251", "koi8-r", "cp866", "iso8859-5", "mac-cyrillic"):
            with self.subTest(encoding=encoding):
                path = self.tmp / "standards" / f"{encoding}.txt"
                path.parent.mkdir(parents=True, exist_ok=True)
                path.write_bytes(source.encode(encoding))
                converted = convert_file(path)
                self.assertEqual(source, converted.text.strip())
                self.assertFalse(converted.warnings, converted.warnings)

    def test_uppercase_title_is_not_mistaken_for_a_wrong_encoding(self):
        # Титульный лист ГОСТа набран прописными целиком. Признак «сплошные
        # прописные» выглядит подозрительно, но здесь он законен.
        source = ("ГОСТ Р 53363-2009 ЦИФРОВЫЕ РАДИОРЕЛЕЙНЫЕ ЛИНИИ СВЯЗИ "
                  "ПОКАЗАТЕЛИ КАЧЕСТВА ТРЕБОВАНИЯ")
        path = self.tmp / "standards" / "титул.txt"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(source.encode("cp1251"))
        converted = convert_file(path)
        self.assertEqual(source, converted.text.strip())
        self.assertEqual("cp1251", converted.meta["encoding"])

    def test_utf16_is_read(self):
        # Блокнот Windows до 2019 года сохранял «Юникод» именно так.
        source = "Заметка о настройке мультиплексора"
        path = self.tmp / "regulations" / "utf16.txt"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"\xff\xfe" + source.encode("utf-16-le"))
        converted = convert_file(path)
        self.assertEqual(source, converted.text.strip())

    def test_unsupported_suffix_gives_warning_not_crash(self):
        # .dwg — чертёж AutoCAD: формат, который система не разбирает и не
        # обещает разбирать. Проверяем, что это предупреждение, а не падение.
        path = self.write("literature/чертёж.dwg", "что-то")
        converted = convert_file(path)
        self.assertEqual(converted.text, "")
        self.assertTrue(converted.warnings)
        self.assertIn("неподдерживаемый формат", converted.warnings[0])

    def test_known_format_with_wrong_content_explains_itself(self):
        # Файл с расширением поддерживаемого формата, но с чужим содержимым:
        # конвертер обязан объяснить, что не так, а не молча вернуть пустоту.
        path = self.write("literature/книга.rtf", "что-то")
        converted = convert_file(path)
        self.assertEqual(converted.text, "")
        self.assertTrue(converted.warnings)
        self.assertIn("RTF", converted.warnings[0])

    def test_missing_file_gives_warning(self):
        converted = convert_file(self.tmp / "нет-такого.md")
        self.assertTrue(converted.is_empty)
        self.assertIn("не найден", converted.warnings[0])


@unittest.skipIf(pymupdf is None, "не установлен pymupdf")
class ConvertPdfTests(TempCase):
    def test_headings_are_restored_by_font_size(self):
        path = self.tmp / "standards" / "методика.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [[
            ("Методика контроля излучения", 20),
            ("1. Условия измерений", 15),
            (LONG_TEXT, 11),
        ]])
        converted = convert_file(path)
        self.assertIn("# Методика контроля излучения", converted.text)
        self.assertIn("## 1. Условия измерений", converted.text)
        self.assertEqual(converted.title, "Методика контроля излучения")
        self.assertEqual(converted.page_count, 1)
        self.assertFalse(converted.needs_ocr)

    def test_hyphenated_words_are_glued(self):
        path = self.tmp / "standards" / "перенос.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [[
            ("Занимаемая полоса измерена в полосе про-", 11),
            ("пускания 30 кГц по методу 99 % мощности.", 11),
        ]])
        converted = convert_file(path)
        self.assertIn("пропускания", converted.text)
        self.assertNotIn("про-", converted.text)

    def test_page_markers_follow_pages(self):
        path = self.tmp / "standards" / "две-страницы.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [
            [("Страница первая", 16), (LONG_TEXT, 11)],
            [("Страница вторая", 16), (LONG_TEXT, 11)],
        ])
        converted = convert_file(path)
        self.assertEqual(converted.page_count, 2)
        self.assertEqual(
            [page for page, _ in convert_module.page_markers(converted.text)], [1, 2]
        )
        self.assertNotIn("<!--", convert_module.strip_page_markers(converted.text))

    def test_empty_pdf_needs_ocr(self):
        path = self.tmp / "literature" / "скан.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [[]])
        converted = convert_file(path)
        self.assertTrue(converted.needs_ocr)
        self.assertTrue(converted.is_empty)
        self.assertTrue(any("OCR" in warning for warning in converted.warnings))

    def test_broken_pdf_does_not_raise(self):
        path = self.tmp / "literature" / "битый.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"%PDF-1.4 \x00\x01 not a pdf")
        converted = convert_file(path)
        self.assertEqual(converted.text, "")
        self.assertTrue(converted.warnings)


@unittest.skipIf(docx is None, "не установлен python-docx")
class ConvertDocxTests(TempCase):
    def setUp(self):
        super().setUp()
        self.path = self.tmp / "reports" / "2024-118.docx"
        self.path.parent.mkdir(parents=True, exist_ok=True)
        make_docx(self.path)
        self.converted = convert_file(self.path)

    def test_headings_from_styles(self):
        self.assertIn("# Отчёт по обращению 2024-118", self.converted.text)
        self.assertIn("## Результаты измерений", self.converted.text)
        self.assertEqual(self.converted.title, "Отчёт по обращению 2024-118")

    def test_table_becomes_markdown_table(self):
        self.assertIn("| Параметр | Значение | Единица |", self.converted.text)
        self.assertIn("| --- | --- | --- |", self.converted.text)
        self.assertIn("| SNR | 13.7 | дБ |", self.converted.text)
        # Вертикальная черта внутри ячейки не должна ломать разметку таблицы.
        self.assertIn("EVM \\| допуск", self.converted.text)

    def test_lists_and_caption(self):
        self.assertIn("- проверить полосовой фильтр", self.converted.text)
        self.assertIn("- заменить кабельную сборку", self.converted.text)
        self.assertIn("Рис. 1 — спектр принимаемого сигнала", self.converted.text)

    def test_broken_docx_does_not_raise(self):
        path = self.tmp / "reports" / "битый.docx"
        path.write_bytes("PK\x03\x04 не документ".encode("utf-8"))
        converted = convert_file(path)
        self.assertTrue(converted.is_empty)
        self.assertTrue(converted.warnings)


# ------------------------------------------------------ хеш и тип документа ---

class HashAndTypeTests(TempCase):
    def test_sha256_matches_hashlib(self):
        path = self.write("standards/a.md", "# Заголовок\n")
        expected = hashlib.sha256(path.read_bytes()).hexdigest()
        self.assertEqual(sha256_file(path), expected)

    def test_sha256_changes_with_content(self):
        path = self.write("standards/a.md", "# Заголовок\n")
        before = sha256_file(path)
        path.write_text("# Другой заголовок\n", encoding="utf-8")
        self.assertNotEqual(before, sha256_file(path))

    def test_guess_doc_type_by_top_directory(self):
        self.assertEqual(guess_doc_type(self.tmp / "standards" / "a.pdf", self.tmp), "standards")
        self.assertEqual(guess_doc_type(self.tmp / "reports" / "x" / "b.md", self.tmp), "reports")

    def test_guess_doc_type_falls_back_to_literature(self):
        self.assertEqual(guess_doc_type(self.tmp / "прочее" / "a.pdf", self.tmp), "literature")
        self.assertEqual(guess_doc_type(self.tmp / "a.pdf", self.tmp), "literature")

    def test_guess_doc_type_without_root(self):
        self.assertEqual(guess_doc_type(self.tmp / "datasheets" / "xyz" / "a.pdf"), "datasheets")


# ------------------------------------------------------------------ чанки ---

class ChunkTests(TempCase):
    def test_matches_corpus_load_file(self):
        source = EXAMPLES / "standards" / "obw-method.md"
        expected = corpus.load_file(source, EXAMPLES)
        converted = convert_file(source)
        produced = chunks_from_markdown(
            converted.text,
            "standards/obw-method",
            "standards",
            {"title": converted.title, "path": "standards/obw-method.md"},
        )
        self.assertEqual(
            [(chunk.chunk_id, chunk.title_path, chunk.text) for chunk in produced],
            [(chunk.chunk_id, chunk.title_path, chunk.text) for chunk in expected],
        )

    def test_meta_contains_title_and_path(self):
        chunks = chunks_from_markdown(
            "# Раздел\n\n" + LONG_TEXT, "literature/книга", "literature",
            {"title": "Книга", "path": "literature/книга.pdf"},
        )
        self.assertTrue(chunks)
        for chunk in chunks:
            self.assertEqual(chunk.meta["title"], "Книга")
            self.assertEqual(chunk.meta["path"], "literature/книга.pdf")
            self.assertEqual(chunk.doc_type, "literature")
            self.assertTrue(chunk.chunk_id.startswith("literature/книга#"))

    def test_page_markers_become_meta_page(self):
        text = (
            f"{convert_module.page_marker(1)}\n\n# Первый раздел\n\n{LONG_TEXT}\n\n"
            f"{convert_module.page_marker(2)}\n\n# Второй раздел\n\n{LONG_TEXT}\n"
        )
        chunks = chunks_from_markdown(text, "standards/док", "standards", {"title": "Док"})
        self.assertEqual([chunk.meta.get("page") for chunk in chunks], [1, 2])
        for chunk in chunks:
            self.assertNotIn("<!--", chunk.text)
        self.assertIn("с. 2", chunks[1].citation)

    def test_front_matter_is_not_indexed(self):
        text = "---\ntitle: Из шапки\nvendor: ACME\n---\n# Раздел\n\n" + LONG_TEXT
        chunks = chunks_from_markdown(text, "literature/книга", "literature")
        self.assertEqual(chunks[0].meta["title"], "Из шапки")
        self.assertEqual(chunks[0].meta["vendor"], "ACME")
        self.assertNotIn("title:", chunks[0].text)


# ------------------------------------------------------------------ приём ---

class IngestFileTests(TempCase):
    def test_adds_document_and_chunks(self):
        path = self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        result = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual((result.added, result.updated, result.skipped, result.failed),
                         (1, 0, 0, 0))
        self.assertGreater(result.chunks, 0)
        self.assertEqual(result.documents, ["standards/обв"])

        document = self.repos.documents.by_doc_id("standards/обв")
        self.assertIsNotNone(document)
        self.assertEqual(document.doc_type, "standards")
        self.assertEqual(document.title, "Методика ОБВ")
        self.assertEqual(document.sha256, sha256_file(path))
        self.assertEqual(document.chunk_count, result.chunks)
        self.assertIsNotNone(document.indexed_at)
        self.assertEqual(self.repos.chunks.count(), result.chunks)

    def test_second_ingest_of_unchanged_file_is_skipped(self):
        path = self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        first = ingest_path(self.repos, path, root=self.tmp)
        second = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual(second.skipped, 1)
        self.assertEqual(second.added, 0)
        self.assertEqual(second.chunks, 0)
        self.assertEqual(second.documents, [])
        self.assertEqual(self.repos.chunks.count(), first.chunks)

    def test_changed_file_is_reindexed(self):
        path = self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        ingest_path(self.repos, path, root=self.tmp)
        path.write_text(
            "# Методика ОБВ, редакция 2\n\n" + LONG_TEXT + "\n\n## Приложение\n\n" + LONG_TEXT,
            encoding="utf-8",
        )
        result = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual((result.added, result.updated), (0, 1))
        document = self.repos.documents.by_doc_id("standards/обв")
        self.assertEqual(document.title, "Методика ОБВ, редакция 2")
        self.assertEqual(document.sha256, sha256_file(path))
        # Чанки заменены целиком, а не добавлены к прежним.
        self.assertEqual(self.repos.chunks.count(), document.chunk_count)

    def test_force_reindexes_unchanged_file(self):
        path = self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        ingest_path(self.repos, path, root=self.tmp)
        result = ingest_path(self.repos, path, root=self.tmp, force=True)
        self.assertEqual((result.updated, result.skipped), (1, 0))
        self.assertGreater(result.chunks, 0)

    def test_explicit_doc_type_and_confidentiality(self):
        path = self.write("прочее/регламент.md", "# Регламент\n\n" + LONG_TEXT)
        ingest_path(self.repos, path, root=self.tmp,
                    doc_type="regulations", confidentiality="nda")
        document = self.repos.documents.by_doc_id("прочее/регламент")
        self.assertEqual(document.doc_type, "regulations")
        self.assertEqual(document.confidentiality, "nda")
        self.assertEqual(self.repos.chunks.all_chunks()[0].doc_type, "regulations")

    def test_broken_file_is_counted_as_failed(self):
        path = self.tmp / "literature" / "битый.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes("%PDF-1.4 \x00 мусор".encode("utf-8"))
        result = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual((result.failed, result.added), (1, 0))
        self.assertTrue(result.warnings)
        self.assertIsNone(self.repos.documents.by_doc_id("literature/битый"))

    def test_missing_file_is_failed_not_exception(self):
        result = ingest_path(self.repos, self.tmp / "нет.md", root=self.tmp)
        self.assertEqual(result.failed, 1)
        self.assertIn("не найден", result.warnings[0])

    def test_remove_document(self):
        path = self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        ingest_path(self.repos, path, root=self.tmp)
        self.assertTrue(remove_document(self.repos, "standards/обв"))
        self.assertIsNone(self.repos.documents.by_doc_id("standards/обв"))
        self.assertEqual(self.repos.chunks.count(), 0)
        self.assertFalse(remove_document(self.repos, "standards/обв"))


@unittest.skipIf(pymupdf is None, "не установлен pymupdf")
class IngestPdfTests(TempCase):
    def test_pdf_chunks_carry_page_numbers(self):
        path = self.tmp / "standards" / "методика.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [
            [("Методика контроля", 18), ("1. Условия", 14), (LONG_TEXT, 11)],
            [("2. Результаты", 14), (LONG_TEXT, 11)],
        ])
        result = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual(result.added, 1)
        chunks = self.repos.chunks.all_chunks()
        self.assertTrue(chunks)
        pages = {chunk.meta.get("page") for chunk in chunks}
        self.assertEqual(pages, {1, 2})
        for chunk in chunks:
            self.assertNotIn("<!--", chunk.text)
            self.assertEqual(chunk.meta["path"], "standards/методика.pdf")
        document = self.repos.documents.by_doc_id("standards/методика")
        self.assertEqual(document.meta["page_count"], 2)

    def test_scanned_pdf_is_reported_not_indexed(self):
        path = self.tmp / "literature" / "скан.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        make_pdf(path, [[]])
        result = ingest_path(self.repos, path, root=self.tmp)
        self.assertEqual(result.failed, 1)
        self.assertTrue(any("OCR" in warning for warning in result.warnings))
        self.assertIsNone(self.repos.documents.by_doc_id("literature/скан"))


class IngestDirectoryTests(TempCase):
    def build_library(self):
        self.write("README.md", "# Служебная заметка\n\n" + LONG_TEXT)
        self.write("standards/обв.md", "# Методика ОБВ\n\n" + LONG_TEXT)
        self.write("reports/2023-041.md", "# Отчёт 2023-041\n\n" + LONG_TEXT)
        self.write("прочее/книга.txt", "Глава 1\n\n" + LONG_TEXT)
        self.write("standards/~$временный.md", "мусор")
        self.write("standards/схема.dwg", "не текст")

    def test_walks_recursively_and_skips_root_files(self):
        self.build_library()
        result = ingest_directory(self.repos, self.tmp)
        self.assertEqual(result.added, 3)
        self.assertEqual(result.failed, 0)
        self.assertIsNone(self.repos.documents.by_doc_id("README"))
        self.assertEqual(
            sorted(document.doc_id for document in self.repos.documents.list()),
            ["reports/2023-041", "standards/обв", "прочее/книга"],
        )

    def test_doc_type_from_top_directory(self):
        self.build_library()
        ingest_directory(self.repos, self.tmp)
        docs = {document.doc_id: document for document in self.repos.documents.list()}
        self.assertEqual(docs["standards/обв"].doc_type, "standards")
        self.assertEqual(docs["reports/2023-041"].doc_type, "reports")
        self.assertEqual(docs["standards/обв"].meta.get("doc_type_source"), "каталог")

    def test_unknown_directory_falls_back_to_content(self):
        """Каталог корпусу не известен — тип берётся из текста.

        «Глава 1» в первой строке — примета книги, и файл «прочее/книга»
        уезжает в literature по делу. А вот список телефонов не относится ни
        к чему: его место на полке «прочее». Раньше полки не было, и такой
        документ молча записывался в книги вместе со всем остальным.
        """
        self.build_library()
        self.write("прочее/телефоны.txt",
                   "Иванов 101\nПетров 102\nСидоров 103\nКузнецов 104\n"
                   "Смирнов 105\nВасильев 106\nПопов 107\nСоколов 108\n")
        ingest_directory(self.repos, self.tmp)
        docs = {document.doc_id: document for document in self.repos.documents.list()}
        self.assertEqual(docs["прочее/книга"].doc_type, "literature")
        phones = docs["прочее/телефоны"]
        self.assertEqual(phones.doc_type, "misc")
        self.assertNotEqual(phones.meta.get("doc_type_source"), "каталог")

    def test_second_run_skips_everything(self):
        self.build_library()
        ingest_directory(self.repos, self.tmp)
        again = ingest_directory(self.repos, self.tmp)
        self.assertEqual((again.added, again.updated, again.skipped), (0, 0, 3))
        forced = ingest_directory(self.repos, self.tmp, force=True)
        self.assertEqual((forced.added, forced.updated, forced.skipped), (0, 3, 0))

    def test_progress_is_reported(self):
        self.build_library()
        messages = []
        ingest_directory(self.repos, self.tmp, progress=messages.append)
        self.assertEqual(len(messages), 3)
        self.assertTrue(all(message.startswith("[") for message in messages))
        self.assertTrue(any("standards/обв.md" in message for message in messages))

    def test_patterns_are_respected(self):
        self.build_library()
        result = ingest_directory(self.repos, self.tmp, patterns=("*.txt",))
        self.assertEqual(result.added, 1)
        self.assertEqual([d.doc_id for d in self.repos.documents.list()], ["прочее/книга"])

    def test_missing_root_raises(self):
        with self.assertRaises(FileNotFoundError):
            ingest_directory(self.repos, self.tmp / "нет-каталога")

    def test_example_corpus_is_ingested(self):
        result = ingest_directory(self.repos, EXAMPLES)
        self.assertEqual(result.added, 5)
        self.assertGreater(result.chunks, 0)
        self.assertEqual(self.repos.chunks.count(), result.chunks)
        found = self.repos.chunks.search_lexical("занимаемая полоса")
        self.assertTrue(found)


class MissingDependencyTests(TempCase):
    """Контур без тяжёлых пакетов: приём обязан объясниться, а не упасть."""

    def test_pdf_without_pymupdf(self):
        path = self.tmp / "standards" / "методика.pdf"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"%PDF-1.4")

        def raise_missing():
            raise convert_module.MissingDependencyError("нужен пакет pymupdf (pip install pymupdf)")

        with mock.patch.object(convert_module, "_import_pymupdf", raise_missing):
            converted = convert_file(path)
            result = ingest_path(self.repos, path, root=self.tmp)
        self.assertTrue(converted.is_empty)
        self.assertIn("pymupdf", converted.warnings[0])
        self.assertEqual(result.failed, 1)

    def test_docx_without_python_docx(self):
        path = self.tmp / "reports" / "отчёт.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"PK")

        def raise_missing():
            raise convert_module.MissingDependencyError("нужен пакет python-docx")

        with mock.patch.object(convert_module, "_import_docx", raise_missing):
            converted = convert_file(path)
        self.assertTrue(converted.is_empty)
        self.assertIn("python-docx", converted.warnings[0])


class SummaryTests(unittest.TestCase):
    def test_summary_is_human_readable(self):
        result = IngestResult(added=2, updated=1, skipped=3, failed=1, chunks=42,
                              warnings=["скан"])
        text = result.summary()
        self.assertIn("добавлено 2", text)
        self.assertIn("обновлено 1", text)
        self.assertIn("без изменений 3", text)
        self.assertIn("с ошибками 1", text)
        self.assertIn("42", text)
        self.assertIn("Предупреждений: 1", text)

    def test_summary_on_empty_run(self):
        self.assertIn("не найдено", IngestResult().summary())

    def test_merge_accumulates(self):
        total = IngestResult()
        total.merge(IngestResult(added=1, chunks=3, documents=["a"]))
        total.merge(IngestResult(skipped=1, warnings=["ой"]))
        self.assertEqual((total.added, total.skipped, total.chunks), (1, 1, 3))
        self.assertEqual(total.documents, ["a"])
        self.assertEqual(total.total, 2)
        self.assertEqual(total.indexed, 1)


if __name__ == "__main__":
    unittest.main()
