"""Тесты экспорта отчёта в DOCX.

Каждый тест собирает документ во временном каталоге и читает его обратно
через python-docx: проверяется не промежуточная структура, а то, что реально
увидит инженер, открывший файл. Сеть не используется — модель подменена
офлайн-заглушкой :class:`reportgen.llm.StubLLM`.
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

import _bootstrap  # noqa: F401
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from reportgen.corpus import load_corpus
from reportgen.export.docx import (
    DRAFT_NOTICE,
    ExportOptions,
    MissingDependencyError,
    export_report,
    footer_for,
    is_draft,
    markdown_to_docx,
)
from reportgen.facts import FactPack
from reportgen.llm import StubLLM
from reportgen.pipeline import Outline, generate_report
from reportgen.retrieval import BM25Index, Retriever

ROOT = Path(__file__).resolve().parents[1]
CASE = ROOT / "examples" / "cases" / "case-2024-118.json"
OUTLINE = ROOT / "templates" / "outline_signal_issue.json"
CORPUS = ROOT / "examples" / "corpus"

MONO_FONT = "Courier New"


def styles_of(document) -> list[str]:
    """Стили абзацев тела документа по порядку."""
    return [paragraph.style.name for paragraph in document.paragraphs]


def texts_of(document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def header_text(document) -> str:
    return "\n".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)


def footer_text(document) -> str:
    return "\n".join(paragraph.text for paragraph in document.sections[0].footer.paragraphs)


def find_run(document, needle: str):
    """Первый run с заданным текстом — по нему проверяется начертание."""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text == needle:
                return run
    raise AssertionError(f"в документе нет фрагмента {needle!r}")


class ExportTestCase(unittest.TestCase):
    """Общая обвязка: временный каталог и сборка документа из строки."""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.addCleanup(self._tmp.cleanup)

    def build(self, markdown: str, name: str = "report.docx", **options):
        """Собирает документ и сразу открывает его обратно."""
        options.setdefault("toc", False)
        options.setdefault("draft", False)
        path = markdown_to_docx(markdown, self.tmp / name, ExportOptions(**options))
        self.assertTrue(path.is_file(), "файл отчёта не создан")
        return Document(str(path))


class MarkdownStructureTests(ExportTestCase):
    """Перенос разметки в документ."""

    def test_heading_levels(self):
        markdown = "\n\n".join(f"{'#' * level} Заголовок {level}" for level in range(1, 7))
        document = self.build(markdown)
        self.assertEqual(styles_of(document), [f"Heading {level}" for level in range(1, 7)])
        self.assertEqual(texts_of(document)[0], "Заголовок 1")

    def test_inline_bold_italic_and_monospace(self):
        document = self.build(
            "Измерение **важное**, значение *ориентировочное*, ключ `snr` в факт-пакете."
        )
        self.assertEqual(len(document.paragraphs), 1)
        self.assertTrue(find_run(document, "важное").bold)
        self.assertTrue(find_run(document, "ориентировочное").italic)
        self.assertEqual(find_run(document, "snr").font.name, MONO_FONT)
        self.assertNotIn("**", document.paragraphs[0].text)
        self.assertNotIn("`", document.paragraphs[0].text)

    def test_escaped_signs_reach_the_document_as_they_are(self):
        """Данные факт-пакета — не разметка.

        Номер группы «*1274*» и модель «Р_168_М» приезжали в документ без
        части знаков: конвертер принимал их за курсив.
        """
        document = self.build("Номер группы: \\*1274\\* и модель Р\\_168\\_М")
        self.assertEqual("Номер группы: *1274* и модель Р_168_М",
                         document.paragraphs[0].text)

    def test_markup_still_works_next_to_escaped_signs(self):
        document = self.build("\\*не курсив\\*, а *курсив* рядом")
        self.assertEqual("*не курсив*, а курсив рядом", document.paragraphs[0].text)
        self.assertTrue(find_run(document, "курсив").italic)

    def test_hard_line_break_keeps_single_paragraph(self):
        # Титульный блок конвейера: строки заканчиваются двумя пробелами.
        document = self.build("**Обращение:** SUP-2024-118  \n**Отправитель:** 1274")
        self.assertEqual(len(document.paragraphs), 1)
        self.assertEqual(
            document.paragraphs[0].text,
            "Обращение: SUP-2024-118\nОтправитель: 1274",
        )

    def test_table_shape_and_bold_header(self):
        markdown = (
            "| Параметр | Значение | Метод |\n"
            "|---|---|---|\n"
            "| ОСШ | 13.7 дБ | по внеполосным участкам |\n"
            "| EVM | 12.4 % | после выравнивания |\n"
        )
        document = self.build(markdown)
        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual((len(table.rows), len(table.columns)), (3, 3))
        self.assertEqual(table.cell(0, 0).text, "Параметр")
        self.assertEqual(table.cell(2, 1).text, "12.4 %")
        self.assertTrue(table.cell(0, 2).paragraphs[0].runs[0].bold)
        self.assertNotEqual(table.cell(1, 0).paragraphs[0].runs[0].bold, True)

    def test_table_alignment_from_separator(self):
        markdown = (
            "| Слева | По центру | Справа |\n"
            "|---|:---:|---:|\n"
            "| a | b | c |\n"
        )
        table = self.build(markdown).tables[0]
        self.assertIsNone(table.cell(1, 0).paragraphs[0].alignment)
        self.assertEqual(table.cell(1, 1).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(table.cell(1, 2).paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)

    def test_table_with_empty_and_ragged_rows(self):
        # python-docx не терпит пустых строк и рваных строк — экспорт их чинит.
        markdown = (
            "| Параметр | Значение |\n"
            "|---|---|\n"
            "|  |  |\n"
            "| Погрешность |  |\n"
            "| Только первый столбец\n"
        )
        table = self.build(markdown).tables[0]
        self.assertEqual((len(table.rows), len(table.columns)), (3, 2))
        self.assertEqual(table.cell(1, 1).text, "")
        self.assertEqual(table.cell(2, 0).text, "Только первый столбец")
        self.assertEqual(table.cell(2, 1).text, "")

    def test_bulleted_and_numbered_lists(self):
        markdown = (
            "- первый\n"
            "- второй\n"
            "  - вложенный\n"
            "\n"
            "1. раз\n"
            "2. два\n"
        )
        document = self.build(markdown)
        self.assertEqual(
            styles_of(document),
            ["List Bullet", "List Bullet", "List Bullet 2", "List Number", "List Number"],
        )
        self.assertEqual(texts_of(document)[2], "вложенный")

    def test_blockquote(self):
        document = self.build("> Статус документа: **ЧЕРНОВИК**.\n> Требует подписи инженера.")
        self.assertEqual(styles_of(document), ["Quote"])
        self.assertEqual(
            document.paragraphs[0].text,
            "Статус документа: ЧЕРНОВИК. Требует подписи инженера.",
        )

    def test_html_comments_are_dropped(self):
        markdown = (
            "<!-- page: 47 -->\n"
            "Видимый абзац.\n"
            "\n"
            "<!-- служебный блок:\n"
            "многострочный комментарий -->\n"
            "\n"
            "Второй абзац <!-- хвост --> продолжается.\n"
        )
        document = self.build(markdown)
        text = "\n".join(texts_of(document))
        self.assertNotIn("page: 47", text)
        self.assertNotIn("служебный блок", text)
        self.assertNotIn("многострочный", text)
        self.assertIn("Видимый абзац.", text)
        self.assertIn("Второй абзац", text)
        self.assertIn("продолжается", text)

    def test_code_block_is_monospaced(self):
        markdown = "Команда:\n\n```bash\n./run.sh verify --report build/report.md\n```\n"
        document = self.build(markdown)
        code = document.paragraphs[-1]
        self.assertEqual(code.text, "./run.sh verify --report build/report.md")
        self.assertEqual(code.runs[0].font.name, MONO_FONT)
        self.assertNotIn("```", "\n".join(texts_of(document)))

    def test_horizontal_rule_becomes_border(self):
        document = self.build("Текст до.\n\n---\n\nТекст после.")
        self.assertEqual(len(document.paragraphs), 3)
        self.assertIn("w:pBdr", document.paragraphs[1]._p.xml)
        self.assertEqual(document.paragraphs[1].text, "")
        self.assertEqual(document.paragraphs[2].text, "Текст после.")


class OptionsTests(ExportTestCase):
    """Колонтитулы, оглавление, приложения, шаблон."""

    def test_draft_notice_in_header_and_page_field_in_footer(self):
        document = self.build(
            "# Отчёт", draft=True, footer_text="Технический отчёт по обращению SUP-2024-118"
        )
        self.assertIn("ЧЕРНОВИК", header_text(document))
        self.assertIn(DRAFT_NOTICE, header_text(document))
        self.assertIn("SUP-2024-118", footer_text(document))
        self.assertIn("PAGE", document.sections[0].footer.paragraphs[0]._p.xml)

    def test_approved_report_has_no_draft_notice(self):
        document = self.build("# Отчёт", draft=False, footer_text="подпись")
        self.assertNotIn("ЧЕРНОВИК", header_text(document))
        self.assertIn("подпись", document.sections[0].footer.paragraphs[0].text)

    def test_toc_field_replaces_manual_contents(self):
        markdown = (
            "# Отчёт\n\n"
            "## Содержание\n\n"
            "1. Исходные данные\n"
            "2. Выводы\n\n"
            "## 1. Исходные данные\n\nтело\n"
        )
        document = self.build(markdown, toc=True)
        self.assertIn("TOC", document.element.xml)
        joined = "\n".join(texts_of(document))
        self.assertIn("Оглавление", joined)
        # Ручной список разделов заменён полем, а не продублирован им.
        self.assertNotIn("2. Выводы", joined)
        self.assertEqual(styles_of(document)[1], "Heading 2")
        self.assertEqual(texts_of(document)[1], "Содержание")

    def test_toc_can_be_disabled(self):
        markdown = (
            "# Отчёт\n\n## Содержание\n\n1. Исходные данные\n\n"
            "## 1. Исходные данные\n\nтело\n"
        )
        document = self.build(markdown, toc=False)
        self.assertNotIn("TOC", document.element.xml)
        self.assertIn("1. Исходные данные", texts_of(document))

    def test_appendix_starts_new_page(self):
        markdown = "# Отчёт\n\n## 1. Выводы\n\nтело\n\n## 2. Приложение А. Источники\n\nтело\n"
        with_break = self.build(markdown, name="a.docx", page_break_before_appendix=True)
        appendix = [p for p in with_break.paragraphs if p.text.startswith("2. Приложение")][0]
        self.assertTrue(appendix.paragraph_format.page_break_before)

        without = self.build(markdown, name="b.docx", page_break_before_appendix=False)
        appendix = [p for p in without.paragraphs if p.text.startswith("2. Приложение")][0]
        self.assertNotEqual(appendix.paragraph_format.page_break_before, True)

    def test_defaults_without_template(self):
        document = self.build("# Отчёт\n\nтело")
        normal = document.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")
        self.assertEqual(normal.font.size, Pt(12))
        self.assertAlmostEqual(normal.paragraph_format.line_spacing, 1.15, places=3)
        section = document.sections[0]
        # Word хранит поля в твипах, поэтому сантиметры сравниваем в них же.
        for margin in (section.top_margin, section.bottom_margin,
                       section.left_margin, section.right_margin):
            self.assertEqual(margin.twips, Cm(2).twips)

    def test_template_provides_styles_margins_and_header(self):
        template = self.tmp / "blank.docx"
        source = Document()
        source.sections[0].left_margin = Cm(3)
        source.styles["Normal"].font.name = "Arial"
        source.sections[0].header.paragraphs[0].text = "АО «Пример», отдел экспертизы"
        source.add_paragraph()  # пустой абзац бланка не должен попасть в отчёт
        source.save(str(template))

        document = self.build("# Отчёт\n\nтело", template=template, draft=True)
        self.assertEqual(document.sections[0].left_margin.twips, Cm(3).twips)
        self.assertEqual(document.styles["Normal"].font.name, "Arial")
        self.assertEqual(texts_of(document)[0], "Отчёт")
        self.assertIn("АО «Пример»", header_text(document))
        self.assertIn(DRAFT_NOTICE, header_text(document))

    def test_missing_template_reports_path(self):
        with self.assertRaises(FileNotFoundError) as caught:
            markdown_to_docx("# Отчёт", self.tmp / "x.docx",
                             ExportOptions(template=self.tmp / "нет-такого.docx"))
        self.assertIn("нет-такого.docx", str(caught.exception))

    def test_creates_parent_directory(self):
        target = self.tmp / "exports" / "2024" / "report.docx"
        result = markdown_to_docx("# Отчёт", target, ExportOptions(toc=False))
        self.assertEqual(result, target)
        self.assertTrue(target.is_file())

    def test_missing_python_docx_gives_russian_message(self):
        saved = sys.modules.get("docx")
        sys.modules["docx"] = None  # имитация контура без python-docx
        try:
            with self.assertRaises(MissingDependencyError) as caught:
                markdown_to_docx("# Отчёт", self.tmp / "нет.docx")
        finally:
            if saved is None:
                sys.modules.pop("docx", None)
            else:
                sys.modules["docx"] = saved
        message = str(caught.exception)
        self.assertIn("python-docx", message)
        self.assertIn("DOCX", message)
        self.assertIsInstance(caught.exception, ImportError)
        self.assertFalse((self.tmp / "нет.docx").exists())


class ExportReportTests(ExportTestCase):
    """Обёртка export_report: статус и обращение попадают в оформление."""

    def test_case_id_goes_to_footer_and_draft_notice_is_on(self):
        path = export_report("# Отчёт\n\nтело", self.tmp / "r.docx", case_id="SUP-2024-118")
        document = Document(str(path))
        self.assertIn("SUP-2024-118", footer_text(document))
        self.assertIn(DRAFT_NOTICE, header_text(document))

    def test_approved_status_removes_draft_notice(self):
        path = export_report("# Отчёт", self.tmp / "r.docx", case_id="SUP-1", status="approved")
        document = Document(str(path))
        self.assertEqual(header_text(document).strip(), "")
        self.assertTrue(is_draft("draft"))
        self.assertTrue(is_draft("verified"))
        self.assertFalse(is_draft("approved"))
        self.assertEqual(footer_for(""), "Технический отчёт")


class RealReportTests(ExportTestCase):
    """Экспорт настоящего отчёта, собранного конвейером на заглушке модели."""

    @classmethod
    def setUpClass(cls) -> None:
        facts = FactPack.load(CASE)
        outline = Outline.load(OUTLINE)
        retriever = Retriever(BM25Index(load_corpus(CORPUS)))
        result = generate_report(
            facts, outline, StubLLM(), retriever, generated_at="2024-07-16"
        )
        cls.markdown = result.markdown
        cls.facts = facts
        cls.outline = outline

    def test_generated_report_exports_and_opens(self):
        path = export_report(
            self.markdown, self.tmp / "SUP-2024-118.docx",
            case_id=self.facts.case_id, status="draft",
        )
        document = Document(str(path))  # файл открывается без исключений

        headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
        self.assertEqual(headings[0], self.outline.title)
        for section in self.outline.sections:
            self.assertTrue(
                any(section.title in heading for heading in headings),
                f"в документе нет раздела «{section.title}»",
            )

        body = "\n".join(texts_of(document))
        self.assertNotIn("<!--", body)
        self.assertNotIn("служебный блок", body)
        self.assertIn("SUP-2024-118", body)

        # Таблицы измерений перенесены целиком, вместе со служебным блоком.
        self.assertGreaterEqual(len(document.tables), len(self.outline.sections) - 1)
        values = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        self.assertIn("13.7 дБ (± 0.4 дБ)", values)
        self.assertIn("facts_digest", values)

        appendix = [p for p in document.paragraphs if p.text.startswith("8. Приложение")]
        self.assertTrue(appendix, "приложение с источниками не найдено")
        self.assertTrue(appendix[0].paragraph_format.page_break_before)

        self.assertIn(self.facts.case_id, footer_text(document))


if __name__ == "__main__":
    unittest.main()
